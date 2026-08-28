# -*- coding: utf-8 -*-
"""
Gera `docs/Artigo_PedroConrado_RichardyRodrigues_ATUALIZADO.docx`: o conteúdo atualizado
do TCC (com as medições já executadas) na mesma apresentação do artigo que foi enviado
ao professor, `docs/Artigo_PedroConrado_RichardyRodrigues.docx`.

Aquele arquivo é usado apenas como ORIGEM (leitura), para duas coisas:
  1. TEMPLATE — dele são herdados styles.xml, docDefaults e a configuração de seção
     (página A4 e margens ABNT). O corpo é esvaziado e reconstruído.
  2. IMAGENS — os 13 diagramas nele renderizados são extraídos e reinseridos com as
     mesmas dimensões; o repositório só guarda o código Mermaid, sem imagens.

Layout herdado da origem:
  - Página A4 (21,0 x 29,7 cm); margens 3 cm esq./sup. e 2 cm dir./inf.
  - Times New Roman 12 pt, justificado, entrelinha simples, 6 pt depois
  - Título e subtítulo centralizados em caixa alta; autoria alinhada à direita

Correções de formatação em relação à origem (defeitos do Pandoc, não do padrão):
  - todos os títulos em 12 pt (a origem usava 16 pt e 14 pt em Heading 1 e 2)
  - seções primárias sempre em caixa alta (a origem tinha 1 e 2 em caixa mista,
    numeradas por lista automática, e 3 a 8 em caixa alta literal)
  - tabelas e figuras com legenda numerada e linha "Fonte:", conforme ABNT
  - tabelas em Times New Roman 10 pt, com fios apenas horizontais (padrão IBGE/ABNT)
    e cabeçalho repetido nas quebras de página

Uso:  python docs/tcc/gerar-artigo-docx.py
"""

import re
import struct
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Emu, Pt

RAIZ = Path(__file__).resolve().parents[2]
ORIGEM = RAIZ / "docs" / "Artigo_PedroConrado_RichardyRodrigues.docx"
SAIDA = RAIZ / "docs" / "Artigo_PedroConrado_RichardyRodrigues_ATUALIZADO.docx"
SAIDA_APENDICE = RAIZ / "docs" / "Apendice_A_Diagramas.docx"
SAIDA_APENDICE_C = RAIZ / "docs" / "Apendice_C_Evidencias.docx"

FONTE = "Times New Roman"
CORPO_PT = Pt(12)
TABELA_PT = Pt(10)
RECUO = Emu(444500)  # recuo de primeira linha (~1,24 cm)

FONTE_PADRAO = "Fonte: elaborado pelos autores (2026)."

# Material suplementar publicado à parte — o artigo não leva apêndice.
# A URL cita uma TAG do git, não uma branch: link de branch muda de conteúdo, e
# referência acadêmica precisa apontar sempre para o mesmo estado do repositório.
REPO_URL = "https://github.com/Richardy-Rodrigues/hubservi"
TAG_TCC = "tcc-v2"
APENDICE_A_URL = f"{REPO_URL}/blob/{TAG_TCC}/docs/tcc/apendice-a-diagramas.md"
APENDICE_B_URL = f"{REPO_URL}/blob/{TAG_TCC}/docs/tcc/apendice-b-reproducao.md"
APENDICE_C_URL = f"{REPO_URL}/blob/{TAG_TCC}/docs/tcc/apendice-c-evidencias.md"

# Links profundos para o registro e para as saídas brutas. O artigo citava estes dois
# como caminhos de sistema de arquivos ("docs/tcc/medicoes/"), o que não é verificável
# por quem lê o PDF: reivindicar reprodutibilidade exige dar ao leitor o endereço.
REGISTRO_URL = f"{REPO_URL}/blob/{TAG_TCC}/docs/tcc/medicoes/registro-medicoes.md"
EVIDENCIAS_URL = f"{REPO_URL}/tree/{TAG_TCC}/docs/tcc/medicoes/evidencias"

PRINTS = RAIZ / "docs" / "tcc" / "medicoes" / "evidencias" / "prints"

# Capturas das execuções das ferramentas, na ordem em que compõem o Apêndice C.
# Esta lista é a ÚNICA fonte da numeração: o apêndice a percorre para gerar as figuras
# C.1 a C.17, e o corpo do artigo consulta `ref_c()` para escrever as remissões. Assim,
# inserir ou remover uma captura não deixa nenhuma referência cruzada desatualizada.
#
# O identificador P-xx é o do catálogo de capturas do repositório; o P-06 (Madge) não
# foi capturado, e a lacuna é declarada na abertura do apêndice em vez de silenciada —
# a medição correspondente (M-03) é atestada pelo arquivo bruto e confirmada por
# segunda ferramenta (dependency-cruiser, P-08).
CAPTURAS = [
    ("P-01", "P-01-reprodutibilidade.png",
     "Ambiente de medição: sistema operacional, processador, Node.js e versões das "
     "ferramentas, no *commit* `ad89e6c`",
     "registro de ambiente", "15 e 16 jul. 2026", "2026-07-15/ambiente.txt"),
    ("P-02", "P-02-vitest.png",
     "Suíte unitária do *baseline*, com os 11 testes existentes no início da avaliação",
     "Vitest 3.2.7", "15 jul. 2026", "2026-07-15/vitest-unit.log"),
    ("P-03", "P-03-cobertura-testes.png",
     "Cobertura após a ampliação da suíte: 31,99% das linhas (559 de 1.747) e 75% dos ramos",
     "Vitest com `@vitest/coverage-v8`", "16 jul. 2026", "2026-07-16/coverage-semana5.txt"),
    ("P-04", "P-04-M-02-eslint.png",
     "Análise estática na configuração do projeto: 19 erros e 9 avisos",
     "ESLint 9.32", "15 jul. 2026", "2026-07-15/eslint-report.json"),
    ("P-05", "P-05-M-21-eslint-sonarjs.png",
     "Análise estática sob a configuração de medição, com regras de complexidade e de "
     "*code smells*: 25 erros e 4 avisos",
     "ESLint 9.32 com `eslint-plugin-sonarjs` 3", "16 jul. 2026",
     "2026-07-16/eslint-sonarjs-report.json"),
    ("P-07", "P-07-M-22-duplicacao-cod-jscpd.png",
     "Duplicação de código no agregado: 3,03% — 10 clones, 107 de 3.528 linhas",
     "jscpd 4", "16 jul. 2026", "2026-07-16/jscpd/jscpd-report.json"),
    ("P-08", "P-08-M-23-24-acoplamento-instabilidade.png",
     "Instabilidade por módulo: núcleo estável (`lib/utils` 4%, `client.ts` 8%, "
     "`AuthContext` 13%) e folhas voláteis entre 90% e 100%",
     "dependency-cruiser 16.10.4", "16 jul. 2026", "2026-07-16/depcruise-metrics.txt"),
    ("P-09", "P-09-M-04-05-stack-local-smoke-integracao.png",
     "Reconstrução do banco a partir do histórico versionado (10 *migrations* e *seed*) e "
     "*smoke* de integração aprovado, 4 de 4",
     "Supabase CLI 2.109.1 e Vitest", "16 jul. 2026", "2026-07-16/vitest-integration.log"),
    ("P-10", "P-10-M-05-13-rls-trigger.png",
     "Suíte completa de autorização (RLS) e de *triggers*, executada contra a API PostgREST real",
     "Vitest (integração)", "16 jul. 2026", "2026-07-16/suite-integracao-completa.log"),
    ("P-11", "P-11-F02-03-antes-correcao.png",
     "**Antes** da correção: duas falhas, com o vazamento literal do campo `email` de outro "
     "usuário na saída do teste",
     "Vitest (integração)", "16 jul. 2026", "2026-07-16/rls-furos-ANTES.log"),
    ("P-12", "P-12-F-02-03-depois-correcao.png",
     "**Depois** da correção: os mesmos dois testes aprovados, após as *migrations* corretivas",
     "Vitest (integração)", "16 jul. 2026", "2026-07-16/rls-furos-DEPOIS.log"),
    ("P-13", "P-13-F-01-privilegio-api-public-profile.png",
     "Privilégios de API concedidos aos papéis `anon`, `authenticated` e `service_role` sobre "
     "`profiles`, após a *migration* que tornou o histórico auto-suficiente",
     "`psql` sobre o *stack* local", "16 jul. 2026", "2026-07-16/grants-profiles-depois.txt"),
    ("P-14", "P-14-M-14-17-lighthouse-antes.png",
     "Carregamento inicial **antes** do *code-splitting*: *performance score* de 85 e LCP de "
     "3,49 s (execução mediana de três)",
     "Lighthouse 12.8.2, perfil móvel", "16 jul. 2026", "2026-07-16/lighthouse-3.json"),
    ("P-15", "P-15-M-14-17-lighthouse-depois.png",
     "Carregamento inicial **depois** do *code-splitting*: *performance score* de 88 e LCP de "
     "3,17 s (execução mediana de três)",
     "Lighthouse 12.8.2, perfil móvel", "16 jul. 2026", "2026-07-16/lighthouse-split-2.json"),
    ("P-16", "P-16-M-18-20-autocannon-carga-listagem.png",
     "Carga sobre a listagem de serviços: 44.413 requisições em 20 s, vazão de 2.221 req/s, "
     "nenhum erro e nenhuma resposta fora da faixa 2xx",
     "autocannon 8", "16 jul. 2026", "2026-07-16/autocannon-load.json"),
    ("P-17", "P-17-M-25-vulnerabilidade-audit.png",
     "Análise de composição de dependências sobre 825 pacotes: 12 vulnerabilidades altas, 8 "
     "moderadas e nenhuma crítica",
     "`npm audit`", "16 jul. 2026", "2026-07-16/npm-audit.json"),
    ("P-18", "P-18-M-26-integridade-schema.png",
     "Verificação de integridade do *schema*: nenhum erro encontrado",
     "`supabase db lint` 2.109.1", "16 jul. 2026", "2026-07-16/supabase-db-lint.txt"),
]

_N_CAPTURA = {pid: i + 1 for i, (pid, *_) in enumerate(CAPTURAS)}
_CAPTURA = {c[0]: c for c in CAPTURAS}


def ref_c(pid):
    """Remissão ao Apêndice C — "Figura C.11" —, calculada a partir de `CAPTURAS`."""
    return f"C.{_N_CAPTURA[pid]}"


def fonte_captura(pid):
    """Linha de fonte de uma captura: ferramenta, data e arquivo bruto preservado."""
    _, _, _, ferramenta, data, arquivo = _CAPTURA[pid]
    return (f"Fonte: {ferramenta}; execução de {data}; saída bruta preservada em "
            f"`evidencias/{arquivo}`.")


_INLINE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*|`(.+?)`", re.DOTALL)


def _largura_png(caminho):
    """Largura em pixels de um PNG, lida do cabeçalho IHDR (sem dependência externa)."""
    with open(caminho, "rb") as f:
        cabecalho = f.read(24)
    if cabecalho[:8] != b"\x89PNG\r\n\x1a\n":
        raise RuntimeError(f"Não é um PNG: {caminho}")
    return struct.unpack(">I", cabecalho[16:20])[0]


def extrair_imagens(caminho):
    """Devolve as imagens de um .docx na ordem em que aparecem no corpo.

    Cada item é (bytes, largura_emu, altura_emu). As dimensões vêm do `wp:extent`
    original, de modo que a reinserção reproduza o mesmo resultado visual.
    """
    doc = Document(str(caminho))
    imagens = []
    for blip in doc.element.body.iter(qn("a:blip")):
        rid = blip.get(qn("r:embed"))
        if rid is None:
            continue
        blob = doc.part.related_parts[rid].blob
        extent = None
        no = blip
        while no is not None and extent is None:
            extent = no.find(qn("wp:extent"))
            no = no.getparent()
        cx = int(extent.get("cx")) if extent is not None else None
        cy = int(extent.get("cy")) if extent is not None else None
        imagens.append((blob, cx, cy))
    return imagens


# ---------------------------------------------------------------------------
# Infraestrutura de layout
# ---------------------------------------------------------------------------

_PRIMARIA = re.compile(r"^\d+\s")


class Artigo:
    """Constrói o documento replicando a formatação do arquivo enviado ao professor."""

    def __init__(self, origem):
        self.doc = Document(str(origem))
        self._esvaziar()
        self.imagens = extrair_imagens(origem)
        self.i_imagem = 0
        self.n_tabela = 0
        self.n_figura = 0
        # Prefixo da numeração de figuras e tabelas. Vazio no artigo ("Figura 1");
        # "A." no apêndice publicado à parte ("Figura A.1"), cuja numeração precisa
        # ser autocontida — inserir uma figura no corpo do artigo não pode deslocá-la.
        self.prefixo = ""

    def pular_imagens(self, n):
        """Avança o cursor de imagens sem inserir nada.

        O apêndice é gerado como documento próprio, mas suas imagens continuam vindo
        da mesma origem, depois das do corpo do artigo. Sem isto, o gerador do apêndice
        começaria pela imagem 1 e trocaria todas as legendas silenciosamente.
        """
        self.i_imagem += n

    def _esvaziar(self):
        corpo = self.doc.element.body
        for filho in list(corpo):
            if filho.tag == qn("w:sectPr"):
                continue  # preserva página, margens e orientação da origem
            corpo.remove(filho)

    # -- runs ---------------------------------------------------------------

    def _run(self, par, texto, bold=False, italic=False, pt=CORPO_PT):
        run = par.add_run(texto)
        run.font.name = FONTE
        run.font.size = pt
        # só marca quando ativo, para que o XML fique idêntico ao do arquivo oficial
        # (lá o corpo herda o não-negrito em vez de desligá-lo explicitamente)
        if bold:
            run.bold = True
        if italic:
            run.italic = True
        # garante a fonte também para caracteres não-ASCII (w:cs / w:eastAsia)
        rpr = run._element.get_or_add_rPr()
        rfonts = rpr.find(qn("w:rFonts"))
        if rfonts is None:
            rfonts = rpr.makeelement(qn("w:rFonts"), {})
            rpr.insert(0, rfonts)
        for attr in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
            rfonts.set(qn(attr), FONTE)
        return run

    def _runs_markdown(self, par, texto, pt=CORPO_PT, bold_base=False, italic_base=False):
        """Converte **negrito**, *itálico* e `código` (este em itálico) em runs.

        A conversão é recursiva, de modo que marcação aninhada — por exemplo
        ``**Snyk para `npm audit`**`` — não deixa delimitadores literais no texto.
        """
        pos = 0
        for m in _INLINE.finditer(texto):
            if m.start() > pos:
                self._run(par, texto[pos:m.start()], bold=bold_base, italic=italic_base, pt=pt)
            negrito, italico, codigo = m.groups()
            if negrito is not None:
                self._runs_markdown(par, negrito, pt=pt, bold_base=True, italic_base=italic_base)
            elif italico is not None:
                self._runs_markdown(par, italico, pt=pt, bold_base=bold_base, italic_base=True)
            else:
                self._run(par, codigo, bold=bold_base, italic=True, pt=pt)
            pos = m.end()
        if pos < len(texto):
            self._run(par, texto[pos:], bold=bold_base, italic=italic_base, pt=pt)

    # -- parágrafos ---------------------------------------------------------

    def centro(self, texto, bold=False):
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._runs_markdown(par, texto, bold_base=bold)
        return par

    def direita(self, texto, bold=False):
        """Bloco de autoria, alinhado à direita como no artigo enviado ao professor."""
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        self._runs_markdown(par, texto, bold_base=bold)
        return par

    def vazio(self):
        return self.doc.add_paragraph()

    def titulo_artigo(self, texto):
        return self.centro(texto.upper(), bold=True)

    def secao(self, texto, primaria=None):
        """Título de seção: esquerda, negrito, 12 pt, sem recuo.

        Seções primárias (`1 Introdução`, `9 Conclusão`, `Referências`…) vão em caixa
        alta, conforme a NBR 6024; subseções, em caixa mista. Isso corrige a
        inconsistência da origem, onde 1 e 2 eram numeradas por lista automática em
        caixa mista e 3 a 8 vinham em caixa alta literal, com corpos de 16 e 14 pt.
        """
        if primaria is None:
            primaria = bool(_PRIMARIA.match(texto))
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._run(par, texto.upper() if primaria else texto, bold=True)
        return par

    def corpo(self, texto):
        """Parágrafo de corpo: justificado, recuo de primeira linha."""
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.first_line_indent = RECUO
        self._runs_markdown(par, texto)
        return par

    def item(self, texto, marcador="• "):
        """Item de lista: justificado, recuado à esquerda, sem recuo de 1ª linha."""
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.left_indent = RECUO
        par.paragraph_format.first_line_indent = Emu(0)
        self._run(par, marcador)
        self._runs_markdown(par, texto)
        return par

    def lista(self, itens):
        for it in itens:
            self.item(it)

    def enumerada(self, itens):
        for i, it in enumerate(itens, 1):
            self.item(it, marcador=f"{i}. ")

    def citacao(self, texto):
        """Enunciado destacado (problema, questão de pesquisa)."""
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        par.paragraph_format.left_indent = Emu(1080000)
        par.paragraph_format.first_line_indent = Emu(0)
        self._runs_markdown(par, texto)
        return par

    def fonte(self, texto=FONTE_PADRAO):
        par = self.doc.add_paragraph()
        par.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._runs_markdown(par, texto, pt=TABELA_PT)
        return par

    # -- tabelas e figuras --------------------------------------------------

    @staticmethod
    def _fio(tag, val):
        e = OxmlElement(f"w:{tag}")
        e.set(qn("w:val"), val)
        e.set(qn("w:sz"), "8")
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), "auto")
        return e

    def _bordas_abnt(self, tbl):
        """Traçado do padrão IBGE/ABNT: só fios horizontais no topo, sob o cabeçalho
        e no rodapé; sem fios verticais nem laterais. É o mesmo desenho das tabelas do
        artigo de origem."""
        bordas = OxmlElement("w:tblBorders")
        for tag, val in (("top", "single"), ("bottom", "single"), ("left", "none"),
                         ("right", "none"), ("insideH", "none"), ("insideV", "none")):
            bordas.append(self._fio(tag, val))
        tbl._tbl.tblPr.append(bordas)

        for celula in tbl.rows[0].cells:
            tc_bordas = OxmlElement("w:tcBorders")
            tc_bordas.append(self._fio("bottom", "single"))
            celula._tc.get_or_add_tcPr().append(tc_bordas)

        # repete o cabeçalho quando a tabela quebra de página
        tr_pr = tbl.rows[0]._tr.get_or_add_trPr()
        cabecalho = OxmlElement("w:tblHeader")
        cabecalho.set(qn("w:val"), "true")
        tr_pr.append(cabecalho)

    def tabela(self, legenda, cabecalho, linhas, fonte=FONTE_PADRAO):
        self.n_tabela += 1
        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.LEFT
        self._runs_markdown(cap, f"Tabela {self.prefixo}{self.n_tabela} — {legenda}", bold_base=True)

        tbl = self.doc.add_table(rows=1, cols=len(cabecalho))
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        tbl.autofit = True

        for celula, texto in zip(tbl.rows[0].cells, cabecalho):
            par = celula.paragraphs[0]
            par.alignment = WD_ALIGN_PARAGRAPH.LEFT
            self._run(par, texto, bold=True, pt=TABELA_PT)

        for linha in linhas:
            celulas = tbl.add_row().cells
            for celula, texto in zip(celulas, linha):
                par = celula.paragraphs[0]
                par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                self._runs_markdown(par, texto, pt=TABELA_PT)

        self._bordas_abnt(tbl)
        if fonte:
            self.fonte(fonte)
        return tbl

    def figura(self, legenda, origem):
        """Insere o próximo diagrama da origem, com legenda numerada e linha de fonte.

        `origem` documenta, no código, de qual arquivo Mermaid a imagem foi gerada; as
        imagens são consumidas na mesma ordem em que aparecem no artigo de origem.
        """
        if self.i_imagem >= len(self.imagens):
            raise RuntimeError(
                f"Figura {self.n_figura + 1} ({origem}) não tem imagem correspondente "
                f"na origem — só há {len(self.imagens)} imagens disponíveis."
            )
        blob, cx, cy = self.imagens[self.i_imagem]
        self.i_imagem += 1
        self.n_figura += 1

        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._runs_markdown(cap, f"Figura {self.prefixo}{self.n_figura} — {legenda}", bold_base=True)

        img = self.doc.add_paragraph()
        img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img.add_run().add_picture(BytesIO(blob), width=Emu(cx), height=Emu(cy))

        rod = self.doc.add_paragraph()
        rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._run(rod, FONTE_PADRAO, pt=TABELA_PT)

    def figura_arquivo(self, legenda, caminho, fonte, largura_cm=15.5):
        """Insere uma captura de tela vinda de arquivo, com legenda numerada e fonte.

        Difere de `figura`: aquela consome os diagramas herdados do .docx de origem, em
        ordem fixa; esta lê um PNG do repositório. As duas compartilham o contador, de
        modo que a numeração do documento continue única e sequencial.

        A linha de fonte não é a genérica "elaborado pelos autores": uma captura de
        execução de ferramenta vale como evidência justamente por dizer qual ferramenta,
        em que data e sobre qual arquivo bruto preservado — por isso `fonte` é
        obrigatória aqui.

        A largura padrão (15,5 cm) é a da mancha de texto A4 com as margens ABNT
        (21 − 3 − 2 cm), com folga; a altura é escalada proporcionalmente. A largura
        efetiva é reduzida quando esticar a captura até lá a deixaria abaixo de 150 dpi:
        um terminal ampliado além da própria resolução fica borrado no papel, e o texto
        da saída — que é o conteúdo da evidência — precisa continuar legível.
        """
        caminho = Path(caminho)
        if not caminho.is_file():
            raise RuntimeError(f"Captura ausente: {caminho}")
        largura_px = _largura_png(caminho)
        largura_cm = min(largura_cm, largura_px / 150 * 2.54)
        self.n_figura += 1

        cap = self.doc.add_paragraph()
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._runs_markdown(cap, f"Figura {self.prefixo}{self.n_figura} — {legenda}",
                            bold_base=True)

        img = self.doc.add_paragraph()
        img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        img.add_run().add_picture(str(caminho), width=Cm(largura_cm))

        rod = self.doc.add_paragraph()
        rod.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._runs_markdown(rod, fonte, pt=TABELA_PT)

    def salvar(self, destino):
        self.doc.save(str(destino))


# ---------------------------------------------------------------------------
# Conteúdo — transposto de docs/tcc/TCC-Hubservi.md
# Números provenientes exclusivamente de docs/tcc/medicoes/registro-medicoes.md
# ---------------------------------------------------------------------------

def capa(a):
    # Título, autoria e vínculos conforme o artigo enviado ao professor.
    a.titulo_artigo(
        "Avaliação Técnica de uma Arquitetura Web baseada em BaaS/Serverless "
        "para Intermediação de Serviços:"
    )
    a.centro("UM ESTUDO DE CASO DA PLATAFORMA HUBSERVI")
    a.vazio()
    a.direita("Pedro Conrado Fernandes Vieira")
    a.direita("Graduando em Engenharia de Software – Uni-FACEF")
    a.direita("Email: opeconrado@gmail.com")
    a.vazio()
    a.direita("Richardy Gabriel Rodrigues da Costa")
    a.direita("Graduando em Engenharia de Software – Uni-FACEF")
    a.direita("Email: richardyrodrigues519@gmail.com")
    a.vazio()
    a.direita("Prof. Daniel Facciolo Pires")
    a.direita("Docente do Departamento de Computação – Uni-FACEF")
    a.direita("Email: daniel@facef.br")
    a.vazio()


def resumo(a):
    a.secao("Resumo")
    a.corpo(
        "Arquiteturas web fundamentadas em *Backend as a Service* (BaaS) e em computação "
        "*serverless* têm sido amplamente adotadas no desenvolvimento de aplicações, por "
        "reduzirem o esforço de implementação e operação da infraestrutura de *backend*. "
        "Entretanto, equipes de desenvolvimento enfrentam dificuldade em validar tecnicamente "
        "se essas arquiteturas atendem aos atributos de qualidade esperados, como segurança, "
        "desempenho, testabilidade e manutenibilidade. Este trabalho tem como objetivo avaliar "
        "tecnicamente a arquitetura de software da plataforma Hubservi — uma aplicação web de "
        "intermediação de serviços construída como *Single Page Application* (SPA) em React e "
        "TypeScript, apoiada pelo BaaS Supabase sobre PostgreSQL —, utilizando métricas e testes "
        "de Engenharia de Software relacionados a segurança, desempenho, testabilidade, "
        "manutenibilidade e confiabilidade, à luz da norma ISO/IEC 25010. A pesquisa "
        "caracteriza-se como aplicada, de abordagem mista, com objetivos exploratórios e "
        "descritivos, conduzida por meio de estudo de caso combinado a experimento técnico. "
        "Como contribuição, propõe-se um procedimento de avaliação arquitetural que articula a "
        "norma ISO/IEC 25010, o método ATAM (*Architecture Tradeoff Analysis Method*) e um "
        "conjunto de ferramentas de teste automatizado, análise estática, desempenho e segurança. "
        "O procedimento foi executado sobre ambiente reprodutível e emitiu vereditos por atributo "
        "e por camada: três defeitos reais de autorização e de integridade foram detectados, "
        "corrigidos e re-medidos; a confiabilidade atende aos critérios definidos; o *backend* "
        "atende ao critério de desempenho, com latência de cauda (p97,5) de 253 ms e 0% de erro "
        "sob carga, ao "
        "passo que o carregamento inicial do *frontend* não o atinge, com LCP de 3,17 s após "
        "otimização, contra a meta de 2,5 s; e a manutenibilidade apresenta estrutura sólida, sem "
        "dependências circulares, com higiene de código abaixo do ideal. A capacidade de localizar "
        "deficiências específicas — e não de atestar qualidade uniforme — evidencia o valor do "
        "procedimento proposto."
    )
    a.corpo(
        "**Palavras-chave:** Arquitetura de software. Avaliação arquitetural. ISO/IEC 25010. "
        "Backend as a Service. Serverless."
    )

    a.secao("Abstract")
    a.corpo(
        "Web architectures based on *Backend as a Service* (BaaS) and *serverless* computing have "
        "been widely adopted in application development, as they reduce the effort of implementing "
        "and operating backend infrastructure. However, development teams struggle to technically "
        "validate whether such architectures meet the expected quality attributes, such as "
        "security, performance, testability, and maintainability. This work aims to technically "
        "evaluate the software architecture of the Hubservi platform — a web application for "
        "service intermediation built as a *Single Page Application* (SPA) using React and "
        "TypeScript, supported by the Supabase BaaS over PostgreSQL —, using Software Engineering "
        "metrics and tests related to security, performance, testability, maintainability, and "
        "reliability, in light of the ISO/IEC 25010 standard. The research is characterized as "
        "applied, with a mixed approach and exploratory and descriptive objectives, conducted "
        "through a case study combined with a technical experiment. As a contribution, it proposes "
        "an architectural evaluation procedure that articulates the ISO/IEC 25010 standard, the "
        "ATAM (*Architecture Tradeoff Analysis Method*), and a set of tools for automated testing, "
        "static analysis, performance, and security. The procedure was executed on a reproducible "
        "environment and issued verdicts per attribute and per layer: three real authorization and "
        "integrity defects were detected, fixed, and re-measured; reliability meets the defined "
        "criteria; the backend meets the performance criterion, with a tail latency (p97.5) of "
        "253 ms and a "
        "0% error rate under load, whereas the initial frontend load does not, with an LCP of "
        "3.17 s after optimization against a 2.5 s target; and maintainability shows a sound "
        "structure, free of circular dependencies, with below-ideal code hygiene. The ability to "
        "pinpoint specific deficiencies — rather than to certify uniform quality — demonstrates "
        "the value of the proposed procedure."
    )
    a.corpo(
        "**Keywords:** Software architecture. Architectural evaluation. ISO/IEC 25010. "
        "Backend as a Service. Serverless."
    )


def secao_1(a):
    a.secao("1 Introdução")

    a.secao("1.1 Contexto")
    a.corpo(
        "A intermediação de serviços — atividade que conecta pessoas que demandam um serviço a "
        "profissionais capazes de executá-lo — tem migrado progressivamente para plataformas "
        "digitais. Aplicações web nesse domínio precisam suportar cadastro e descoberta de "
        "serviços, solicitação e gerenciamento de atendimentos e avaliação reputacional, com "
        "requisitos não triviais de segurança, desempenho e confiabilidade, uma vez que manipulam "
        "dados pessoais e transações entre partes que, em geral, não se conhecem previamente."
    )
    a.corpo(
        "Em paralelo, o modelo de desenvolvimento de aplicações web tem sido reconfigurado pela "
        "popularização de arquiteturas baseadas em *Backend as a Service* (BaaS) e em computação "
        "*serverless*. Nesses modelos, responsabilidades tradicionalmente implementadas em "
        "servidores próprios — autenticação, persistência, autorização e regras de acesso — são "
        "delegadas a serviços gerenciados por terceiros, acessados diretamente pelo cliente por "
        "meio de APIs e bibliotecas. Plataformas como o Supabase materializam esse paradigma ao "
        "expor, sobre um banco PostgreSQL, autenticação integrada e autorização declarativa via "
        "*Row Level Security* (RLS), eliminando boa parte da camada de servidor de aplicação "
        "convencional."
    )
    a.corpo(
        "A plataforma Hubservi, objeto deste estudo, é uma aplicação web de intermediação de "
        "serviços construída como *Single Page Application* (SPA) em React e TypeScript, apoiada "
        "pelo Supabase. Trata-se, portanto, de um caso representativo da arquitetura **SPA + BaaS "
        "+ Serverless**, na qual a lógica de negócio se distribui entre o cliente (validações, "
        "fluxos de interface e orquestração de chamadas) e o banco de dados (regras declarativas "
        "de autorização, *triggers* e *views*)."
    )

    a.secao("1.2 Problema de pesquisa")
    a.corpo(
        "A adoção de arquiteturas BaaS/Serverless reduz o esforço de implementação e operação de "
        "infraestrutura, mas desloca atributos de qualidade críticos — em especial a segurança e a "
        "confiabilidade — para configurações declarativas (políticas de RLS, *triggers*, "
        "restrições de integridade) e para a fronteira cliente–serviço gerenciado. Esse "
        "deslocamento dificulta a verificação de que a arquitetura efetivamente atende aos "
        "atributos de qualidade esperados: a ausência de uma camada de servidor de aplicação "
        "própria altera o que deve ser testado, onde residem os pontos de falha e como o "
        "desempenho e a manutenibilidade devem ser medidos."
    )
    a.corpo("Formaliza-se, assim, o problema de pesquisa:")
    a.citacao(
        "Equipes de desenvolvimento de aplicações web têm dificuldade em validar tecnicamente se "
        "uma arquitetura baseada em *Backend as a Service* (BaaS) e *Serverless* atende aos "
        "atributos de qualidade esperados, como segurança, desempenho, testabilidade e "
        "manutenibilidade."
    )

    a.secao("1.3 Questão de pesquisa")
    a.corpo("Decorre do problema a seguinte questão de pesquisa:")
    a.citacao(
        "Como avaliar tecnicamente uma arquitetura web baseada em BaaS/Serverless por meio de "
        "métricas e testes de Engenharia de Software, considerando atributos de qualidade em uma "
        "plataforma de intermediação de serviços?"
    )

    a.secao("1.4 Objetivos")
    a.secao("1.4.1 Objetivo geral")
    a.corpo(
        "Avaliar tecnicamente a arquitetura de software da plataforma Hubservi, baseada em React, "
        "TypeScript, Supabase e PostgreSQL, utilizando métricas e testes relacionados à segurança, "
        "ao desempenho, à testabilidade e à manutenibilidade."
    )
    a.secao("1.4.2 Objetivos específicos")
    a.enumerada([
        "Identificar os requisitos arquiteturais e os atributos de qualidade relevantes para a "
        "plataforma.",
        "Modelar a arquitetura da plataforma.",
        "Documentar os componentes, fluxos, persistência e regras de negócio.",
        "Definir cenários de avaliação arquitetural.",
        "Executar testes automatizados.",
        "Executar análise estática de código.",
        "Executar testes de segurança.",
        "Executar testes de desempenho.",
        "Analisar os resultados obtidos.",
    ])
    a.corpo(
        "**Nota de escopo.** Os objetivos específicos 1 a 4 (delimitação, modelagem, documentação "
        "e definição de cenários) e os objetivos 5 a 9 (execução de testes automatizados, análise "
        "estática, testes de segurança e de desempenho, e análise) foram executados; os resultados "
        "são reportados na Seção 7 e derivam do registro reprodutível em `docs/tcc/medicoes/`."
    )

    a.secao("1.5 Delimitação do escopo")
    a.corpo(
        "**Objeto de pesquisa e instrumento.** O objeto deste trabalho é a **avaliação técnica** "
        "de uma arquitetura web baseada em BaaS/Serverless; a plataforma Hubservi é o "
        "**instrumento** por meio do qual essa avaliação se torna possível. A distinção é "
        "necessária porque o trabalho envolve dois produtos de naturezas distintas: um sistema de "
        "software, desenvolvido integralmente pela equipe, e um procedimento de avaliação, "
        "aplicado sobre esse sistema. Apenas o segundo constitui a contribuição científica "
        "pretendida. Conforme exposto na Seção 1.6, a lacuna identificada é de ordem metodológica "
        "— não é evidente *como* aplicar métricas e testes de Engenharia de Software para validar "
        "atributos de qualidade em um arranjo no qual autorização e integridade migram para o "
        "banco de dados. O que se oferece como contribuição, portanto, é um roteiro de avaliação "
        "reprodutível e transferível a outras aplicações que adotem o mesmo paradigma, e não a "
        "plataforma em si."
    )
    a.corpo(
        "**Justificativa do desenvolvimento do sistema.** A construção do Hubservi não constitui "
        "objetivo do trabalho, mas **pré-requisito metodológico** dele. Avaliar políticas de RLS, "
        "*triggers*, restrições de integridade e a fronteira entre cliente e serviço gerenciado "
        "exige acesso irrestrito ao código-fonte, ao esquema do banco de dados, às configurações "
        "de autorização e ao ambiente de execução — condições inviáveis de obter sobre uma "
        "plataforma de terceiros, cujo código e cuja base de dados não são acessíveis ao "
        "pesquisador. O desenvolvimento próprio é o que assegura a validade interna do "
        "experimento: permite controlar as variáveis do ambiente, reproduzir as medições sobre um "
        "estado conhecido e conduzir o ciclo de detecção e correção de defeitos sem restrições de "
        "acesso. Os quatro fluxos críticos implementados — autenticação, cadastro e busca de "
        "serviços, contratação (*booking*) e avaliação (*review*), descritos na Seção 4 — "
        "constituem a superfície sobre a qual os cenários de avaliação da Seção 5 são exercitados."
    )
    a.corpo(
        "**Delimitação negativa.** Não integram o escopo deste trabalho: (i) a avaliação de "
        "**usabilidade e acessibilidade**, por estarem fora das quatro características da ISO/IEC "
        "25010 (2011) selecionadas como relevantes para o problema de pesquisa — segurança, "
        "eficiência de desempenho, manutenibilidade e confiabilidade —, ainda que constem como "
        "requisitos não funcionais do produto; (ii) o **módulo de recomendação**, que permanece "
        "secundário e limitado à ordenação de resultados por popularidade e avaliação média, não "
        "sendo objeto de avaliação; (iii) arquiteturas de **microsserviços**, uma vez que o "
        "sistema avaliado adota o modelo SPA + BaaS + Serverless, sem camada de servidor de "
        "aplicação própria; e (iv) testes **fim a fim conduzidos por ferramenta dedicada de "
        "automação de navegador**, dado que, em uma arquitetura BaaS, a superfície de integração "
        "relevante é a própria API do serviço gerenciado — os fluxos ponta a ponta são, por essa "
        "razão, verificados por testes de integração executados contra a API, conforme a Seção 5."
    )

    a.secao("1.6 Justificativa")
    a.corpo(
        "A literatura de arquitetura de software dispõe de métodos consolidados de avaliação — "
        "notadamente o ATAM (*Architecture Tradeoff Analysis Method*), de Clements, Kazman e Klein "
        "(2002) — e de modelos de qualidade reconhecidos, como o estabelecido pela norma ISO/IEC "
        "25010 (2011). Tais referenciais, contudo, foram concebidos predominantemente sob o "
        "pressuposto de arquiteturas com camada de servidor de aplicação explícita. A crescente "
        "adoção de arquiteturas BaaS/Serverless, nas quais responsabilidades de autorização e "
        "integridade migram para o banco de dados e para serviços gerenciados, cria uma lacuna "
        "prática: não é evidente *como* aplicar métricas e testes de Engenharia de Software para "
        "validar tecnicamente esses atributos de qualidade nesse novo arranjo."
    )
    a.corpo(
        "Justifica-se, portanto, conduzir um estudo de caso que (i) modele e documente "
        "rigorosamente uma arquitetura BaaS/Serverless real, (ii) defina cenários e métricas de "
        "avaliação alinhados aos atributos de qualidade da ISO/IEC 25010, e (iii) organize um "
        "procedimento reprodutível de avaliação. A contribuição é tanto prática — para a equipe da "
        "plataforma Hubservi — quanto metodológica, ao oferecer um roteiro de avaliação técnica "
        "transferível a outras aplicações que adotem o mesmo paradigma."
    )

    a.secao("1.7 Organização do artigo")
    a.corpo(
        "O restante do artigo está organizado da seguinte forma. A Seção 2 apresenta o referencial "
        "teórico sobre arquitetura de software, avaliação arquitetural, qualidade de software e "
        "arquiteturas BaaS/Serverless. A Seção 3 descreve a metodologia adotada. A Seção 4 "
        "documenta a arquitetura da plataforma Hubservi. A Seção 5 detalha o planejamento "
        "experimental e o plano de métricas. A Seção 6 apresenta a avaliação arquitetural com base "
        "no ATAM. A Seção 7 consolida os resultados das medições executadas. A Seção 8 apresenta a "
        "conclusão, as limitações e os trabalhos futuros e, por fim, são listadas as referências."
    )

    a.secao("1.8 Disponibilidade dos artefatos")
    a.corpo(
        "Todo o material que sustenta este trabalho é público e versionado, em coerência com a "
        "própria tese do artigo: uma avaliação técnica só é verificável se o instrumento, o "
        "procedimento e as evidências puderem ser inspecionados por terceiros. O repositório "
        f"({REPO_URL}) reúne o código-fonte da plataforma, as *migrations* que definem o esquema e "
        "as políticas de autorização, as suítes de teste, o registro completo das medições — com "
        "ferramenta, versão, data, valor, critério e veredito — e as saídas brutas de cada execução."
    )
    a.corpo(
        "Três documentos são publicados como **material suplementar**, referenciados ao longo do "
        "texto. O primeiro (VIEIRA; COSTA, 2026a) reúne os diagramas de modelagem — UML, BPMN e DER "
        "— e o dicionário de dados, numerados de forma autocontida (Figuras A.1 a A.10 e Tabelas "
        "A.1 a A.7). O segundo (VIEIRA; COSTA, 2026b) descreve como reproduzir cada medição "
        "reportada na Seção 7, classificando-as pelo que exigem do ambiente e explicitando quais "
        "**não** são reproduzíveis e por quê — informação que a Seção 5.5 trata como parte das "
        "ameaças à validade, e não como omissão. O terceiro (VIEIRA; COSTA, 2026c) reúne as "
        "capturas de tela de todas as execuções de ferramenta que sustentam a Seção 7, das quais "
        "as mais relevantes são reproduzidas no próprio corpo do artigo."
    )


def secao_2(a):
    a.secao("2 Referencial Teórico")
    a.corpo(
        "Esta seção fundamenta teoricamente o trabalho em quatro eixos: (i) arquitetura de "
        "software e atributos de qualidade; (ii) avaliação arquitetural e o método ATAM; (iii) "
        "qualidade de software segundo a ISO/IEC 25010; e (iv) o paradigma de arquiteturas "
        "BaaS/Serverless. Encerra com a contribuição da Engenharia de Software no que tange a "
        "testes e análise estática."
    )

    a.secao("2.1 Arquitetura de software e atributos de qualidade")
    a.corpo(
        "A arquitetura de software de um sistema é definida por Bass, Clements e Kazman (2012) "
        "como o conjunto de estruturas necessárias para raciocinar sobre o sistema, compreendendo "
        "elementos de software, as relações entre eles e as propriedades de ambos. Para os "
        "autores, a arquitetura é o artefato que viabiliza ou inibe os atributos de qualidade do "
        "sistema: decisões arquiteturais — e não primariamente decisões de implementação — "
        "determinam o grau em que requisitos como desempenho, segurança e modificabilidade serão "
        "satisfeitos."
    )
    a.corpo(
        "Bass, Clements e Kazman (2012) distinguem requisitos funcionais, que expressam *o que* o "
        "sistema deve fazer, dos *requisitos de atributos de qualidade*, que expressam *quão bem* "
        "o sistema deve fazê-lo. Esses requisitos são expressos por meio de **cenários de "
        "atributos de qualidade**, estruturados em seis partes: fonte do estímulo, estímulo, "
        "artefato, ambiente, resposta e medida da resposta. Tal estrutura é central para este "
        "trabalho, pois fornece o formato pelo qual os atributos avaliados na plataforma Hubservi "
        "são operacionalizados como cenários verificáveis (Seções 5 e 6)."
    )
    a.corpo(
        "Os autores também introduzem o conceito de **táticas** — decisões de projeto que "
        "influenciam o controle de um atributo de qualidade — e de **ASR** (*Architecturally "
        "Significant Requirements*), os requisitos cuja satisfação depende de decisões "
        "arquiteturais. No contexto de uma arquitetura BaaS/Serverless, táticas de segurança como "
        "*autenticar atores* e *autorizar atores* materializam-se em mecanismos declarativos "
        "(autenticação gerenciada e políticas de RLS), o que justifica avaliá-las de forma "
        "específica."
    )

    a.secao("2.2 Avaliação arquitetural e o método ATAM")
    a.corpo(
        "Clements, Kazman e Klein (2002), em *Evaluating Software Architectures*, argumentam que a "
        "arquitetura, por ser o primeiro artefato em que os atributos de qualidade do sistema se "
        "tornam analisáveis, pode e deve ser avaliada antes de a construção avançar, reduzindo o "
        "risco de retrabalho. Os autores propõem o **ATAM** (*Architecture Tradeoff Analysis "
        "Method*), método de avaliação baseado em cenários cujo objetivo não é fornecer notas "
        "precisas, mas identificar **riscos**, **pontos de sensibilidade** (*sensitivity points*) "
        "e **pontos de compromisso** (*tradeoff points*) decorrentes das decisões arquiteturais."
    )
    a.corpo(
        "O ATAM organiza-se em torno de uma **árvore de utilidade** (*utility tree*), que decompõe "
        "a utilidade geral do sistema em atributos de qualidade, estes em refinamentos, e estes, "
        "por fim, em cenários priorizados segundo a importância para o negócio e o grau de risco "
        "arquitetural. Os principais conceitos do método, mobilizados na Seção 6, são:"
    )
    a.lista([
        "**Ponto de sensibilidade:** propriedade de um ou mais componentes da arquitetura que é "
        "crítica para se alcançar uma resposta de atributo de qualidade.",
        "**Ponto de compromisso:** propriedade que é ponto de sensibilidade para mais de um "
        "atributo, de modo que melhorá-la para um atributo pode degradar outro.",
        "**Risco e não risco:** decisões arquiteturais com consequências potencialmente negativas "
        "(ou explicitamente seguras) para os atributos de qualidade.",
    ])
    a.corpo(
        "O método é especialmente adequado a este trabalho por ser orientado a cenários e por "
        "focalizar *tradeoffs*, dimensão central em arquiteturas BaaS/Serverless, nas quais a "
        "delegação de responsabilidades a serviços gerenciados implica compromissos explícitos "
        "entre simplicidade, controle, desempenho e segurança."
    )

    a.secao("2.3 Qualidade de software e a norma ISO/IEC 25010")
    a.corpo(
        "A norma ISO/IEC 25010 (2011), parte da família SQuaRE (*Systems and software Quality "
        "Requirements and Evaluation*), define um **modelo de qualidade do produto de software** "
        "composto por oito características, cada qual subdividida em subcaracterísticas. As "
        "características são: adequação funcional, eficiência de desempenho, compatibilidade, "
        "usabilidade, confiabilidade, segurança, manutenibilidade e portabilidade."
    )
    a.corpo("Para o escopo deste trabalho, são mobilizadas as seguintes características e subcaracterísticas:")
    a.lista([
        "**Segurança** (*security*): confidencialidade, integridade, não repúdio, "
        "responsabilização (*accountability*) e autenticidade. Avaliada por meio de autenticação, "
        "autorização e controle de acesso indevido.",
        "**Eficiência de desempenho** (*performance efficiency*): comportamento temporal, "
        "utilização de recursos e capacidade. Avaliada por tempo de resposta, tempo de "
        "carregamento e comportamento sob carga.",
        "**Manutenibilidade** (*maintainability*): modularidade, reusabilidade, analisabilidade, "
        "modificabilidade e **testabilidade**. Cabe registrar que, na ISO/IEC 25010 (2011), a "
        "testabilidade é uma subcaracterística da manutenibilidade; neste trabalho ela é tratada "
        "como dimensão avaliativa destacada, em razão de sua centralidade para a verificação dos "
        "demais atributos.",
        "**Confiabilidade** (*reliability*): maturidade, disponibilidade, tolerância a falhas e "
        "recuperabilidade. Avaliada pela consistência de operações e pelo comportamento em fluxos "
        "críticos.",
    ])
    a.corpo(
        "A norma fornece, assim, o vocabulário e a taxonomia que estruturam o plano de métricas "
        "(Seção 5), garantindo rastreabilidade entre cada métrica coletada e a característica de "
        "qualidade que ela pretende evidenciar."
    )

    a.secao("2.4 Arquiteturas BaaS e Serverless")
    a.corpo(
        "O termo *serverless* designa um modelo de execução no qual a provisão, o escalonamento e "
        "a manutenção de servidores são abstraídos e delegados a um provedor, de modo que a equipe "
        "de desenvolvimento concentra-se na lógica da aplicação. Uma de suas manifestações é o "
        "*Backend as a Service* (BaaS), no qual funcionalidades de *backend* comumente necessárias "
        "— autenticação, banco de dados, armazenamento de arquivos e autorização — são oferecidas "
        "como serviços gerenciados, consumidos diretamente pelo cliente por meio de SDKs e APIs."
    )
    a.corpo(
        "Nesse arranjo, parte significativa das regras de negócio e, sobretudo, das regras de "
        "autorização desloca-se para a camada de dados. No caso do Supabase, isso se concretiza "
        "por meio do *Row Level Security* (RLS) do PostgreSQL, mecanismo que permite definir, de "
        "forma **declarativa**, políticas que restringem quais linhas cada usuário pode ler ou "
        "modificar, avaliadas pelo próprio banco a cada operação. Complementam o modelo as *views* "
        "(que expõem projeções controladas dos dados) e os *triggers* (que aplicam regras de "
        "integridade e máquinas de estado no servidor)."
    )
    a.corpo(
        "Do ponto de vista arquitetural, esse paradigma apresenta implicações relevantes para a "
        "avaliação de qualidade:"
    )
    a.lista([
        "A **superfície de autorização** concentra-se em políticas declarativas, cuja correção "
        "precisa ser testada explicitamente, pois falhas de RLS podem expor dados sem que haja "
        "erro funcional aparente no cliente.",
        "A **ausência de servidor de aplicação próprio** transfere parte do desempenho percebido "
        "para o cliente (carregamento da SPA) e para a latência das chamadas ao serviço gerenciado.",
        "A **testabilidade** passa a depender da capacidade de testar tanto o cliente quanto as "
        "regras residentes no banco.",
    ])
    a.corpo(
        "Essas implicações motivam a necessidade, identificada no problema de pesquisa, de um "
        "procedimento de avaliação técnica específico para arquiteturas BaaS/Serverless."
    )

    a.secao("2.5 Engenharia de software: testes e análise estática")
    a.corpo(
        "Pressman e Maxim (2016) sistematizam o teste de software como atividade planejada e "
        "mensurável, distinguindo níveis (unidade, integração, sistema) e abordagens (caixa-branca "
        "e caixa-preta), e enfatizam o papel das métricas de software na avaliação objetiva da "
        "qualidade de produto e de processo. Sommerville (2011), por sua vez, situa a verificação "
        "e a validação, a análise estática e a inspeção de código como práticas complementares ao "
        "teste dinâmico, destacando que a análise estática permite detectar classes de defeitos — "
        "e indicadores de manutenibilidade, como complexidade e duplicação — sem a execução do "
        "programa."
    )
    a.corpo(
        "Esses referenciais embasam a escolha das ferramentas e métricas do plano experimental "
        "(Seção 5): testes automatizados (de unidade e de integração) para adequação funcional, "
        "confiabilidade e testabilidade; análise estática para manutenibilidade; e técnicas "
        "específicas para os atributos de segurança e desempenho."
    )


def secao_3(a):
    a.secao("3 Metodologia")

    a.secao("3.1 Classificação da pesquisa")
    a.corpo(
        "A pesquisa é classificada segundo quatro dimensões usuais na metodologia científica: "
        "natureza, abordagem, objetivos e procedimentos técnicos."
    )
    a.tabela(
        "Classificação da pesquisa",
        ["Dimensão", "Classificação", "Justificativa"],
        [
            ["**Natureza**", "Aplicada",
             "Visa gerar conhecimento de aplicação prática — um procedimento de avaliação técnica "
             "— dirigido à solução de um problema concreto de validação arquitetural."],
            ["**Abordagem**", "Mista (quantitativa e qualitativa)",
             "Combina a coleta de métricas objetivas (cobertura, tempos de resposta, complexidade, "
             "vulnerabilidades) com a análise qualitativa de decisões arquiteturais via ATAM."],
            ["**Objetivos**", "Exploratória e descritiva",
             "Explora a aplicação de métodos de avaliação a um paradigma arquitetural pouco "
             "coberto pela literatura clássica (BaaS/Serverless) e descreve detalhadamente a "
             "arquitetura e os resultados da avaliação."],
            ["**Procedimentos**", "Estudo de caso com experimento técnico",
             "Investiga em profundidade um caso real (Hubservi) e conduz um experimento técnico "
             "controlado de coleta de métricas e execução de testes."],
        ],
    )

    a.secao("3.2 Objeto de estudo")
    a.corpo(
        "O objeto de estudo é a plataforma **Hubservi**, aplicação web de intermediação de "
        "serviços construída como SPA em React e TypeScript e apoiada pelo BaaS Supabase sobre "
        "PostgreSQL. A escolha justifica-se por ser um caso representativo do paradigma SPA + BaaS "
        "+ Serverless, com regras de negócio expressivas residentes no banco (políticas de RLS, "
        "*triggers* e *views*), o que a torna adequada à investigação do problema de pesquisa. A "
        "arquitetura do objeto é detalhada na Seção 4."
    )

    a.secao("3.2.1 Construção do instrumento")
    a.corpo(
        "Conforme a delimitação da Seção 1.5, a plataforma é o **instrumento** da avaliação, "
        "desenvolvido integralmente pela equipe. Sua construção seguiu um processo estruturado, "
        "guiado por artefatos de Engenharia de Software, e não é objeto de avaliação em si — "
        "registra-se aqui por dois motivos: a **validade interna** do experimento depende de o "
        "instrumento ser conhecido e controlado, e parte desses artefatos alimenta diretamente as "
        "etapas seguintes da pesquisa."
    )
    a.lista([
        "**Iniciação e concepção:** Termo de Abertura do Projeto (TAP), Modelo Canvas e análise "
        "SWOT delimitaram problema, escopo (MVP), premissas, restrições e critérios de sucesso.",
        "**Planejamento:** Estrutura Analítica do Projeto (EAP) e 5W2H organizaram as entregas; os "
        "requisitos foram especificados e priorizados sob identificadores rastreáveis (**RF-xx** "
        "funcionais, **RNF-xx** não funcionais).",
        "**Modelagem:** UML (casos de uso, classes, componentes, sequência e implantação), BPMN "
        "(contratação e gerenciamento de *booking*) e modelo de dados (DER e dicionário de dados).",
        "**Construção incremental:** entrega por marcos — **M1** requisitos e arquitetura; **M2** "
        "fluxos centrais (autenticação, serviços, *booking*); **M3** segurança e governança de "
        "dados (RLS e *triggers*); **M4** validação de qualidade; **M5** evolução da descoberta de "
        "serviços (ordenação e filtros).",
        "**Fluxo de qualidade contínuo:** *lint*, testes e revisão a cada alteração, conforme "
        "premissa registrada no TAP.",
    ])
    a.corpo(
        "Os artefatos de modelagem, além de documentarem o instrumento, **atendem aos objetivos "
        "específicos 2 e 3** e constituem a base a partir da qual se derivam a árvore de utilidade "
        "do ATAM (Seção 6) e os cenários de avaliação (Seção 5) — ou seja, ligam a construção do "
        "instrumento à avaliação propriamente dita. Encontram-se reproduzidos no material "
        "suplementar (VIEIRA; COSTA, 2026a)."
    )

    a.secao("3.3 Etapas metodológicas")
    a.corpo(
        "A condução do trabalho organiza-se em seis etapas, alinhadas aos objetivos específicos "
        "(Seção 1.4.2):"
    )
    a.enumerada([
        "**Levantamento de requisitos arquiteturais e atributos de qualidade.** Identificação dos "
        "*Architecturally Significant Requirements* (ASR) e seleção das características de "
        "qualidade da ISO/IEC 25010 (2011) relevantes ao domínio: segurança, eficiência de "
        "desempenho, manutenibilidade (com ênfase em testabilidade) e confiabilidade.",
        "**Modelagem da arquitetura.** Recuperação e representação da arquitetura real a partir do "
        "código-fonte e das *migrations*, produzindo modelos UML (casos de uso, classes, "
        "sequência, componentes e implantação), modelo de processos de negócio (BPMN) e o modelo "
        "de dados (DER e dicionário de dados). Os artefatos constam do material suplementar "
        "(VIEIRA; COSTA, 2026a).",
        "**Documentação técnica.** Descrição dos componentes, fluxos, persistência e regras de "
        "negócio (Seção 4), assegurando rastreabilidade entre cada afirmação e sua fonte no "
        "repositório.",
        "**Definição de cenários de avaliação.** Construção de uma árvore de utilidade ATAM e "
        "especificação de cenários de atributo de qualidade no formato de seis partes (Seção 6), "
        "além do mapeamento atributo, cenário, métrica, critério e ferramenta (Seção 5).",
        "**Execução da avaliação técnica.** Execução de testes automatizados (unitários, de "
        "componente e de integração contra a API), análise estática de código, testes de segurança "
        "e testes de desempenho, conforme o plano de métricas. O ambiente de teste de banco/API "
        "foi levantado com uma instância Supabase local, garantindo reprodutibilidade e isolamento "
        "em relação a dados de produção.",
        "**Análise dos resultados.** Interpretação das métricas coletadas à luz dos critérios "
        "definidos e dos *tradeoffs* identificados no ATAM, com discussão dos achados e "
        "recomendações (Seção 7).",
    ])

    a.secao("3.4 Instrumentos e ferramentas")
    a.corpo(
        "A coleta de dados apoia-se em um conjunto de ferramentas, organizado por atributo de "
        "qualidade e detalhado na Seção 5. Onde a ferramenta inicialmente prevista não estava "
        "disponível no ambiente (por exigir conta ou instalação indisponível), adotou-se um "
        "substituto **da mesma classe**, medindo as mesmas métricas; cada substituição está "
        f"registrada no registro de medições ({REGISTRO_URL})."
    )
    a.lista([
        "**Testes automatizados:** Vitest e React Testing Library (unitários e de componente); "
        "Vitest contra o *stack* Supabase local via PostgREST (integração).",
        "**Análise estática:** ESLint; `eslint-plugin-sonarjs` para *code smells* e complexidade — "
        "em substituição ao SonarQube, que exige serviço externo; `jscpd` para duplicação.",
        "**Análise de dependências e modularização:** Madge e dependency-cruiser.",
        "**Desempenho:** Lighthouse (carregamento inicial); autocannon (carga sobre a API) — em "
        "substituição ao k6/JMeter.",
        "**Segurança:** `npm audit` para análise de composição de dependências (SCA) — em "
        "substituição ao Snyk, que exige conta; `supabase db lint` para o *schema*. A varredura "
        "dinâmica (DAST/OWASP ZAP) permanece pendente de execução contra a URL de produção.",
    ])

    a.secao("3.5 Procedimentos de análise")
    a.corpo(
        "As métricas quantitativas são confrontadas com os critérios de aceitação previamente "
        "definidos (limiares e metas), e os achados qualitativos organizados segundo os conceitos "
        "do ATAM (riscos, pontos de sensibilidade e pontos de compromisso). A triangulação entre "
        "as evidências quantitativas e a análise arquitetural qualitativa sustenta a resposta à "
        "questão de pesquisa."
    )
    a.corpo(
        "Adotou-se, na execução, o ciclo **medir, detectar, corrigir e re-medir**: quando uma "
        "medição reprova o sistema, o defeito é corrigido e a medição repetida, reportando-se o "
        "par antes/depois. Um conjunto de cenários que aprovasse o sistema já na primeira execução "
        "não demonstraria qualidade, e sim fraqueza dos cenários — por isso os dois estados são "
        "preservados como evidência."
    )
    a.corpo(
        "**Regra de registro e anti-fabricação.** Toda medição registra ferramenta e versão, "
        "ambiente, configuração, data, valor obtido e veredito, com o artefato bruto "
        f"correspondente arquivado publicamente ({EVIDENCIAS_URL}) e reproduzido como captura de "
        "tela no Apêndice C (VIEIRA; COSTA, 2026c). Nenhum valor é reportado sem evidência "
        "associada, e os resultados desfavoráveis são registrados com o mesmo rigor que os "
        "favoráveis (Seção 7)."
    )


def secao_4(a):
    a.secao("4 Arquitetura do Hubservi")
    a.corpo(
        "Esta seção documenta a arquitetura real da plataforma Hubservi, recuperada a partir do "
        "código-fonte (`src/`) e das *migrations* do banco de dados (`supabase/migrations/`). A "
        "descrição responde aos objetivos específicos 2 e 3 (modelar e documentar componentes, "
        "fluxos, persistência e regras de negócio)."
    )

    a.secao("4.1 Visão geral e estilo arquitetural")
    a.corpo(
        "O Hubservi adota o estilo **SPA + BaaS + Serverless**. Um *Single Page Application* (SPA) "
        "executado no navegador concentra a interface e a orquestração da lógica de aplicação e "
        "comunica-se diretamente com o BaaS Supabase, que provê autenticação, banco de dados "
        "PostgreSQL e autorização declarativa via *Row Level Security* (RLS). Não há, portanto, "
        "servidor de aplicação intermediário desenvolvido sob medida, tampouco decomposição em "
        "microsserviços: as responsabilidades de *backend* são delegadas a serviços gerenciados, e "
        "parte expressiva das regras de negócio reside no próprio banco, sob a forma de "
        "*triggers*, *views* e políticas de RLS."
    )
    a.corpo("A arquitetura organiza-se em camadas lógicas, conforme a Figura 1.")
    a.figura(
        "Camadas lógicas da arquitetura do Hubservi",
        "docs/tcc/TCC-Hubservi.md, Seção 4.1 (diagrama Mermaid)",
    )
    a.tabela(
        "Camadas lógicas e elementos correspondentes no repositório",
        ["Camada", "Responsabilidade", "Elementos no repositório"],
        [
            ["Apresentação", "Renderização, navegação e interação",
             "`src/pages/`, `src/components/`, `src/components/ui/`"],
            ["Aplicação", "Sessão, estado de servidor, regras de fluxo e validação",
             "`src/contexts/AuthContext.tsx`, React Query, esquemas Zod"],
            ["Integração", "Acesso ao BaaS",
             "`src/integrations/supabase/client.ts`, `views.ts`, `types.ts`"],
            ["Dados", "Persistência e regras declarativas",
             "`supabase/migrations/` (PostgreSQL)"],
        ],
    )
    a.corpo(
        "Os diagramas detalhados de componentes e de implantação constam do material suplementar "
        "(VIEIRA; COSTA, 2026a, Figuras A.7 e A.8)."
    )

    a.secao("4.2 Tecnologias")
    a.tabela(
        "Tecnologias empregadas na plataforma Hubservi",
        ["Camada", "Tecnologia", "Versão"],
        [
            ["Biblioteca de UI", "React", "18.3.1"],
            ["Linguagem", "TypeScript", "5.8.3"],
            ["Empacotador / *build*", "Vite", "5.4.19"],
            ["Roteamento", "React Router DOM", "6.30.1"],
            ["Estado de servidor", "TanStack React Query", "5.83.0"],
            ["Formulários e validação", "react-hook-form + Zod", "7.61 / 3.25"],
            ["Componentes de UI", "shadcn/ui (Radix UI) + Tailwind CSS", "3.4.17"],
            ["Cliente BaaS", "@supabase/supabase-js", "2.98.0"],
            ["Banco de dados", "PostgreSQL (gerenciado pelo Supabase)", "—"],
        ],
    )

    a.secao("4.3 Componentes e rotas")
    a.corpo("A aplicação define as rotas apresentadas na Tabela 4 (`src/App.tsx`).")
    a.tabela(
        "Rotas da aplicação e respectivo nível de acesso",
        ["Rota", "Página", "Acesso"],
        [
            ["`/`", "`Index.tsx` — *landing page*", "Público"],
            ["`/auth`", "`Auth.tsx` — login e cadastro", "Público"],
            ["`/services`", "`Services.tsx` — busca e listagem", "Público"],
            ["`/services/:id`", "`ServiceDetail.tsx` — detalhe, avaliações e solicitação", "Público"],
            ["`/dashboard`", "`Dashboard.tsx` — painel por perfil", "Protegido"],
            ["`*`", "`NotFound.tsx` — 404", "Público"],
        ],
    )
    a.corpo(
        "O acesso à rota protegida é controlado pelo componente `ProtectedRoute.tsx`, que "
        "redireciona usuários sem sessão para `/auth`. O `Dashboard.tsx` seleciona a interface "
        "conforme o tipo de usuário: `ClientDashboard` (cliente) ou `ProviderDashboard` "
        "(prestador). Provedores globais configurados na raiz incluem `QueryClientProvider` (cache "
        "e sincronização de dados), `AuthProvider` (sessão e perfil), `BrowserRouter` e provedores "
        "de *feedback* de interface. Os diagramas de casos de uso e de classes constam do material "
        "suplementar (VIEIRA; COSTA, 2026a, Figuras A.2 e A.3)."
    )

    a.secao("4.4 Modelo de dados")
    a.corpo("O esquema do banco compreende cinco tabelas, três tipos enumerados e duas *views*.")
    a.corpo("**Tabelas:** `profiles`, `categories`, `services`, `bookings`, `reviews`.")
    a.corpo("**Enumerações:**")
    a.lista([
        "`user_type`: `client`, `provider`;",
        "`price_type`: `fixed`, `hourly`, `negotiable`;",
        "`booking_status`: `pending`, `accepted`, `completed`, `rejected`, `cancelled`.",
    ])
    a.corpo("**Views:**")
    a.lista([
        "`service_stats` — agrega, por serviço, a contagem de avaliações (`review_count`) e a "
        "média de notas (`average_rating`); definida com `security_invoker = true`.",
        "`public_profiles` — projeção de `profiles` sem dados pessoais sensíveis (expõe `id`, "
        "`full_name`, `avatar_url`, `user_type`, `created_at`; **omite** `email` e `phone`), "
        "utilizada para apresentar dados de perfil a usuários não autenticados.",
    ])
    a.corpo(
        "As chaves estrangeiras estabelecem as relações: `services` referencia `profiles` "
        "(prestador) e `categories`; `bookings` referencia `services` e duas vezes `profiles` "
        "(cliente e prestador); `reviews` referencia `services` e `profiles` (cliente e "
        "prestador). A tabela `reviews` possui a restrição de unicidade (`service_id`, "
        "`client_id`), garantindo no máximo uma avaliação por cliente por serviço. O DER e o "
        "dicionário de dados constam do material suplementar (VIEIRA; COSTA, 2026a, Figura A.1 e "
        "Tabelas A.1 a A.7)."
    )

    a.secao("4.5 Fluxos principais")
    a.secao("4.5.1 Autenticação")
    a.corpo(
        "O cadastro (`Auth.tsx`) chama `supabase.auth.signUp`, fornecendo `full_name` e "
        "`user_type` em metadados. O *trigger* `on_auth_user_created` materializa, de forma "
        "idempotente, a linha correspondente em `profiles`. O `AuthContext` assina "
        "`onAuthStateChange` e carrega o perfil do usuário autenticado (VIEIRA; COSTA, 2026a, "
        "Figura A.4)."
    )
    a.secao("4.5.2 Cadastro e busca de serviços")
    a.corpo(
        "Prestadores criam, editam e removem serviços pelo `ServiceForm` no `ProviderDashboard`, "
        "com validação Zod. A busca (`Services.tsx`) consulta serviços ativos (`is_active = "
        "true`), com filtro por categoria e por título, paginação e ordenação por **recência**, "
        "**avaliação** ou **popularidade** — estas duas últimas apoiadas na *view* "
        "`service_stats`. Cabe destacar que a recomendação de serviços, no Hubservi, resume-se a "
        "essa ordenação por popularidade e avaliação, constituindo módulo secundário e não o foco "
        "deste trabalho."
    )
    a.secao("4.5.3 Contratação (booking)")
    a.corpo(
        "A solicitação (`BookingDialog`) insere um registro em `bookings` com `status = "
        "'pending'`. O cliente acompanha e pode cancelar solicitações pendentes; o prestador "
        "aceita, rejeita, conclui ou cancela (VIEIRA; COSTA, 2026a, Figuras A.5 e A.9)."
    )
    a.secao("4.5.4 Avaliação (review)")
    a.corpo(
        "Concluído um *booking*, o cliente pode registrar uma avaliação (nota de 1 a 5 e "
        "comentário) pelo `ReviewForm` (VIEIRA; COSTA, 2026a, Figura A.6)."
    )

    a.secao("4.6 Regras de negócio residentes no banco")
    a.corpo(
        "Boa parte das invariantes do domínio é imposta no servidor por *triggers* e restrições, e "
        "não apenas no cliente — característica marcante do paradigma BaaS."
    )
    a.tabela(
        "Regras de negócio implementadas no banco de dados",
        ["Regra", "Mecanismo", "Fonte"],
        [
            ["Provisão automática de perfil no cadastro",
             "*trigger* `on_auth_user_created` para `handle_new_user()` (idempotente, "
             "`SECURITY DEFINER`)", "*migration* inicial; `20260316201000`"],
            ["`updated_at` atualizado a cada alteração",
             "*trigger* `update_updated_at_column()`", "*migration* inicial"],
            ["Máquina de estados do *booking*",
             "*trigger* `validate_booking_status_transition()`",
             "*migration* inicial; `20260528000000`"],
            ["Imutabilidade do `user_type` (anti-escalonamento de privilégio)",
             "*trigger* `prevent_user_type_change()`", "`20260514100000`"],
            ["`booking.provider_id` deve coincidir com o dono do serviço",
             "*trigger* `validate_booking_provider()`", "`20260514100100`"],
            ["`price_max` maior ou igual a `price_min`",
             "*constraint* `services_price_range_check`", "`20260514100300`"],
            ["Avaliação só após *booking* `completed`",
             "política RLS de `INSERT` em `reviews`", "`20260514100400`"],
            ["Cliente pode cancelar *booking* pendente",
             "política RLS e transição `pending` para `cancelled`", "`20260528000000`"],
            ["`review.provider_id` deve coincidir com o dono do serviço",
             "*trigger* de validação espelhado (correção reportada em 7.3.2)",
             "*migration* de correção (julho de 2026)"],
        ],
    )
    a.corpo(
        "A máquina de estados do *booking* admite as transições `pending` para `accepted`, "
        "`rejected` ou `cancelled`, e `accepted` para `completed` ou `cancelled`; transições "
        "inválidas resultam em exceção (VIEIRA; COSTA, 2026a, Figura A.10)."
    )

    a.secao("4.7 Segurança: autorização declarativa via RLS")
    a.corpo(
        "Todas as tabelas têm RLS habilitado. A autorização é expressa por políticas declarativas "
        "avaliadas pelo PostgreSQL a cada operação. A Tabela 6 resume as políticas vigentes, já "
        "**no estado posterior às correções** de segurança reportadas na Seção 7.3.2."
    )
    a.tabela(
        "Políticas de Row Level Security por tabela e operação",
        ["Tabela", "Operação", "Política (resumo)"],
        [
            ["`profiles`", "SELECT/UPDATE/INSERT",
             "usuário acessa e edita apenas o próprio perfil (`auth.uid() = id`); dados de "
             "contraparte são obtidos pela *view* `public_profiles`, sem PII"],
            ["`categories`", "SELECT", "leitura pública"],
            ["`services`", "SELECT",
             "qualquer um vê serviços ativos; prestador vê os próprios (ativos ou não)"],
            ["`services`", "INSERT/UPDATE/DELETE",
             "apenas o prestador dono (`auth.uid() = provider_id`)"],
            ["`bookings`", "SELECT", "cliente vê os próprios; prestador vê os próprios"],
            ["`bookings`", "INSERT", "apenas o cliente (`auth.uid() = client_id`)"],
            ["`bookings`", "UPDATE",
             "prestador altera status; cliente pode cancelar apenas os próprios pendentes"],
            ["`reviews`", "SELECT", "leitura pública"],
            ["`reviews`", "INSERT",
             "cliente com *booking* `completed` no serviço; `provider_id` validado por *trigger*"],
            ["`reviews`", "UPDATE/DELETE", "apenas o autor (`auth.uid() = client_id`)"],
        ],
    )
    a.corpo(
        "A proteção de dados pessoais (e-mail e telefone) é assegurada pela restrição da leitura "
        "direta de `profiles` ao próprio registro, combinada à *view* `public_profiles` para "
        "consumo por terceiros e por usuários anônimos. Registre-se que essa configuração é "
        "resultado da avaliação: a política originalmente vigente admitia leitura por qualquer "
        "usuário autenticado (`USING (auth.uid() IS NOT NULL)`), defeito detectado e corrigido "
        "conforme a Seção 7.3.2. Esse arranjo — autorização concentrada em políticas declarativas "
        "— é precisamente o ponto que a avaliação de segurança (Seções 5 e 6) exercitou de forma "
        "sistemática."
    )

    a.secao("4.8 Síntese arquitetural")
    a.corpo(
        "O Hubservi exemplifica os *tradeoffs* característicos do paradigma BaaS/Serverless: "
        "ganha-se simplicidade operacional e velocidade de desenvolvimento ao delegar "
        "autenticação, persistência e autorização a serviços gerenciados, ao custo de concentrar a "
        "correção da segurança em configurações declarativas e de transferir parte do desempenho "
        "percebido para o cliente e para a latência das chamadas ao serviço. Esses pontos "
        "orientaram o planejamento experimental apresentado a seguir."
    )


def secao_5(a):
    a.secao("5 Planejamento Experimental")
    a.corpo(
        "Esta seção operacionaliza a avaliação técnica, atendendo aos objetivos específicos 4 a 8. "
        "Para cada atributo de qualidade da ISO/IEC 25010 (2011) considerado relevante, definiu-se "
        "um conjunto de cenários, métricas, critérios de aceitação e ferramentas. Os critérios "
        "aqui apresentados são **metas e limiares de planejamento**, fixados antes da coleta; os "
        "valores efetivamente medidos constam da Seção 7."
    )

    a.secao("5.1 Visão geral do procedimento de medição")
    a.corpo(
        "O procedimento encadeia atributo de qualidade, cenário em seis partes, métrica, critério "
        "de aceitação, ferramenta, coleta e registro, e análise em face do critério, conforme a "
        "Figura 2."
    )
    a.figura(
        "Procedimento de medição adotado",
        "docs/tcc/TCC-Hubservi.md, Seção 5.1 (diagrama Mermaid)",
    )
    a.corpo(
        "Cada execução de medição registrou: ferramenta e versão, ambiente, configuração, data, "
        "valor obtido e veredito (atende ou não atende ao critério). A reprodutibilidade foi "
        "assegurada pela fixação de ambiente e parâmetros, conforme a etapa 5 da metodologia "
        "(Seção 3.3)."
    )

    a.secao("5.2 Mapeamento entre atributo, cenário, métrica, critério e ferramenta")
    a.corpo(
        "Os valores da coluna **Critério (meta)** são alvos de planejamento; não representam "
        "resultados medidos. A coluna **Ferramenta** já registra os substitutos efetivamente "
        "utilizados, conforme a Seção 3.4."
    )

    a.secao("5.2.1 Segurança")
    a.tabela(
        "Cenários, métricas e critérios de segurança",
        ["Cenário", "Métrica", "Critério (meta)", "Ferramenta"],
        [
            ["Usuário tenta ler ou alterar dados de outro usuário (*booking*, perfil, serviço)",
             "Nº de acessos indevidos bem-sucedidos", "0 acessos indevidos",
             "Testes de política RLS contra a API"],
            ["Cliente tenta avaliar serviço sem *booking* concluído",
             "Nº de inserções de *review* indevidas", "0 inserções",
             "Teste de integração contra a API"],
            ["Usuário tenta alterar o próprio `user_type` (escalonamento)",
             "Tentativa bloqueada (sim/não)", "100% bloqueadas",
             "Teste de integração e verificação do *trigger*"],
            ["Varredura de vulnerabilidades da aplicação web",
             "Nº de vulnerabilidades por severidade", "0 de severidade alta ou crítica",
             "OWASP ZAP (pendente — ver 7.3.2)"],
            ["Vulnerabilidades em dependências", "Nº de CVEs por severidade",
             "0 de severidade alta ou crítica não tratadas",
             "`npm audit` (em substituição ao Snyk)"],
            ["Exposição de PII a usuário anônimo **e a usuário autenticado não relacionado**",
             "Campos sensíveis expostos (`email`, `phone`)",
             "0 campos expostos a ambos os perfis de acesso",
             "Teste de política RLS e inspeção da *view* `public_profiles`"],
            ["Conformidade do *schema* do banco", "Nº de erros de *lint* de *schema*",
             "0 erros", "`supabase db lint`"],
        ],
    )
    a.corpo(
        "**Nota de refinamento do cenário de exposição de PII.** Em sua formulação inicial, este "
        "cenário considerava apenas o acesso **anônimo**. A inspeção das políticas vigentes "
        "evidenciou que a política de leitura de `profiles` admitia qualquer usuário autenticado "
        "(`USING (auth.uid() IS NOT NULL)`), de modo que o cenário restrito ao anônimo seria "
        "satisfeito sem exercitar o perfil de acesso mais permissivo efetivamente existente. O "
        "cenário foi, por isso, ampliado para contemplar também o **usuário autenticado não "
        "relacionado ao perfil consultado**. Registra-se o ajuste por transparência metodológica: "
        "a formulação de cenários é atividade sujeita a refinamento à luz da arquitetura concreta "
        "sob avaliação, e um cenário incapaz de reprovar o sistema não constitui instrumento de "
        "medição."
    )

    a.secao("5.2.2 Eficiência de desempenho")
    a.tabela(
        "Cenários, métricas e critérios de eficiência de desempenho",
        ["Cenário", "Métrica", "Critério (meta)", "Ferramenta"],
        [
            ["Carregamento inicial da SPA", "*Performance score*; LCP; TBT; CLS",
             "Score maior ou igual a 90; LCP menor ou igual a 2,5 s", "Lighthouse"],
            ["Listagem e busca de serviços sob carga",
             "Tempo de resposta (p95); vazão (req/s); taxa de erro",
             "p95 menor ou igual a 800 ms; erro abaixo de 1%",
             "autocannon (em substituição ao k6)"],
            ["Operações de *booking* sob carga concorrente", "Latência (p95); taxa de erro",
             "p95 menor ou igual a 800 ms; erro abaixo de 1%",
             "autocannon (em substituição ao k6/JMeter)"],
            ["Comportamento sob carga sustentada",
             "Estabilidade de latência e de erros ao longo do tempo",
             "Sem degradação progressiva", "autocannon (em substituição ao JMeter)"],
        ],
    )
    a.corpo(
        "O limiar de 800 ms para a latência p95, em aberto no planejamento inicial, foi calibrado "
        "antes da coleta definitiva a partir de uma execução exploratória de baixa intensidade (5 "
        "conexões por 5 s, latência de cauda aproximada de 57 ms) e da heurística de 1 s para a "
        "percepção de fluidez pelo usuário, adotando-se margem conservadora."
    )
    a.corpo(
        "**Percentil efetivamente reportado.** A ferramenta empregada (autocannon) não expõe o "
        "p95 no conjunto padrão de percentis de seu histograma — os valores vizinhos são o p90 e "
        "o p97,5. Os resultados de latência sob carga apresentados na Seção 7 são, portanto, de "
        "**p97,5**, percentil mais exigente que o p95 fixado como critério: satisfeito o limiar "
        "em p97,5, ele está necessariamente satisfeito em p95. Registra-se a diferença para que o "
        "número publicado não prometa precisão superior à que o instrumento entrega."
    )

    a.secao("5.2.3 Testabilidade (subcaracterística de manutenibilidade)")
    a.tabela(
        "Cenários, métricas e critérios de testabilidade",
        ["Cenário", "Métrica", "Critério (meta)", "Ferramenta"],
        [
            ["Cobertura da suíte de testes", "Cobertura de linhas e de ramos (%)",
             "Piso global anti-regressão de 30% e cobertura maior ou igual a 75% nos módulos "
             "críticos (limiar definido na execução — ver 7.3.4)", "Vitest (*coverage* v8)"],
            ["Isolamento de componentes em teste",
             "Nº de componentes testáveis sem dependências externas reais (uso de *mocks*)",
             "Fluxos críticos cobertos por teste isolado", "Vitest e React Testing Library"],
            ["Esforço de criação de teste por fluxo crítico",
             "Existência de teste para cada fluxo crítico (autenticação, *booking*, *review*)",
             "100% dos fluxos críticos com ao menos um teste", "Vitest e React Testing Library"],
        ],
    )

    a.secao("5.2.4 Manutenibilidade")
    a.tabela(
        "Cenários, métricas e critérios de manutenibilidade",
        ["Cenário", "Métrica", "Critério (meta)", "Ferramenta"],
        [
            ["Conformidade de estilo e antipadrões", "Nº de violações de *lint*",
             "Tendência a 0; sem erros", "ESLint"],
            ["Complexidade e *code smells*",
             "Complexidade ciclomática; densidade de *code smells*",
             "Dentro dos limiares das regras adotadas",
             "`eslint-plugin-sonarjs` (em substituição ao SonarQube)"],
            ["Duplicação de código", "% de linhas duplicadas",
             "Menor ou igual a 3%", "`jscpd` (em substituição ao SonarQube)"],
            ["Modularização e acoplamento",
             "Nº de dependências circulares; grafo de dependências", "0 ciclos",
             "Madge e dependency-cruiser"],
        ],
    )

    a.secao("5.2.5 Confiabilidade")
    a.tabela(
        "Cenários, métricas e critérios de confiabilidade",
        ["Cenário", "Métrica", "Critério (meta)", "Ferramenta"],
        [
            ["Consistência da máquina de estados do *booking*",
             "Nº de transições inválidas aceitas", "0 transições inválidas",
             "Teste de integração do *trigger*"],
            ["Integridade referencial em exclusões em cascata",
             "Nº de registros órfãos após exclusão", "0 órfãos", "Teste de integração"],
            ["Fluxos críticos ponta a ponta (autenticação, *booking*, *review*)",
             "Taxa de sucesso dos fluxos", "100% nos casos válidos", "Testes de integração"],
        ],
    )

    a.secao("5.3 Ferramentas: papel e disponibilidade")
    a.corpo(
        "A Tabela 12 consolida as ferramentas efetivamente empregadas na execução, registrando as "
        "substituições em relação ao plano inicial. Todas as trocas preservaram a classe de "
        "ferramenta e as métricas coletadas, conforme a Seção 3.4."
    )
    a.tabela(
        "Ferramentas empregadas e substituições em relação ao plano inicial",
        ["Categoria", "Ferramenta empregada", "Prevista no plano inicial", "Situação"],
        [
            ["Teste automatizado", "Vitest, React Testing Library", "Vitest, React Testing Library",
             "Utilizada; suíte ampliada na execução"],
            ["Teste de integração", "Vitest contra Supabase local via PostgREST",
             "Testes contra a API", "Utilizada"],
            ["Análise estática", "ESLint (*flat config*, ESLint 9)", "ESLint",
             "Utilizada; complementada com regras de qualidade"],
            ["Complexidade e *code smells*", "`eslint-plugin-sonarjs`", "SonarQube",
             "Substituída (SonarQube exige serviço externo)"],
            ["Duplicação", "`jscpd`", "SonarQube", "Substituída"],
            ["Dependências e modularização", "Madge, dependency-cruiser",
             "Madge, dependency-cruiser", "Utilizadas"],
            ["Desempenho (carregamento)", "Lighthouse", "Lighthouse",
             "Utilizada sobre o *build* de produção"],
            ["Desempenho (carga)", "autocannon", "k6, JMeter", "Substituída"],
            ["Segurança (SCA)", "`npm audit`", "Snyk", "Substituída (Snyk exige conta)"],
            ["Segurança (*schema*)", "`supabase db lint`", "—", "Acrescentada na execução"],
            ["Segurança (DAST)", "OWASP ZAP", "OWASP ZAP",
             "**Pendente** — depende da URL de produção"],
        ],
    )

    a.secao("5.4 Ambiente experimental")
    a.corpo(
        "Uma medição só é reprodutível se o ambiente em que ocorreu for conhecido. A Tabela 13 "
        "caracteriza o ambiente único em que toda a coleta foi realizada, nas datas de 15 e 16 de "
        "julho de 2026, sobre o estado do repositório fixado pelo *commit* `ad89e6c`. Os testes de "
        "desempenho de carregamento utilizaram o *build* de produção (`vite build`); as medições "
        "de segurança, de autorização (RLS) e de confiabilidade ocorreram contra uma instância "
        "Supabase local em contêiner, reconstruída do zero a cada execução por `supabase db "
        "reset`, o que elimina qualquer contato com dados reais de produção. Os valores da tabela "
        f"são transcritos dos arquivos `ambiente.txt` arquivados com as evidências ({EVIDENCIAS_URL})."
    )
    a.tabela(
        "Caracterização do ambiente experimental",
        ["Dimensão", "Configuração registrada"],
        [
            ["*Host*",
             "Microsoft Windows 11 Pro for Workstations, *build* 26200; processador AMD Ryzen 5 "
             "PRO 230; 15,2 GB de memória RAM"],
            ["Tempo de execução",
             "Node.js 24.14.0; npm 11.9.0 (versão fixada em `packageManager`); dependências "
             "instaladas por `npm ci`, a partir do `package-lock.json`"],
            ["Aplicação avaliada",
             "*Build* de produção gerado por Vite 5.4.19; servido localmente para as execuções do "
             "Lighthouse"],
            ["Banco de dados e API",
             "Supabase CLI 2.109.1, invocada por `npx` com versão fixada; imagem de contêiner "
             "`supabase/postgres:15.8.1.085` (PostgreSQL 15.8), executada em Docker Desktop sobre "
             "WSL 2; API PostgREST exposta em `127.0.0.1:54321`"],
            ["Versões das ferramentas (Tabela 12)",
             "Vitest 3.2.7 com `@vitest/coverage-v8`; ESLint 9.32 com `typescript-eslint` 8.38 e "
             "`eslint-plugin-sonarjs` 3; jscpd 4; Madge 8; dependency-cruiser 16.10.4; Lighthouse "
             "12.8.2 (perfil móvel, com limitação de CPU e de rede); autocannon 8; `npm audit`; "
             "`supabase db lint` 2.109.1"],
            ["Volume de dados",
             "Base reconstruída do zero (10 *migrations* e *seed*), com 11 categorias e nenhum "
             "usuário pré-existente; as suítes de integração criam e descartam os próprios "
             "usuários e registros a cada execução; o teste de carga foi semeado com 50 serviços "
             "ativos de um prestador, consultados com `limit=20` por requisição"],
        ],
        fonte="Fonte: elaborado pelos autores (2026), a partir dos registros de ambiente das "
              "coletas de 15 e 16 de julho de 2026.",
    )
    a.corpo(
        "Uma dimensão do ambiente **não foi registrada** na coleta: a versão do Docker Desktop. "
        "Registra-se a omissão em vez de preenchê-la com a versão instalada hoje, que não seria a "
        "da medição. O impacto é considerado baixo, porque o que determina o comportamento do "
        "banco é a imagem de contêiner, essa sim fixada por *tag* (`15.8.1.085`), e não o "
        "programa que a executa."
    )

    a.secao("5.5 Ameaças à validade")
    a.lista([
        "**Ambiente local em vez de produção (validade externa).** As medições de API e de banco "
        "correm contra a instância local, sem latência de rede real nem os limites do plano de "
        "serviço gerenciado. Os números de desempenho de *backend* devem ser lidos como um "
        "**piso** — uma execução contra a instância gerenciada tende a apresentar latência maior.",
        "**Volume de dados reduzido (validade de construção).** O teste de carga exercita uma "
        "tabela com 50 serviços, muito abaixo de qualquer operação real. O resultado atesta que a "
        "camada de dados responde sob concorrência, não que se mantenha sob volume; latência de "
        "leitura é sensível ao tamanho da relação e à seletividade dos índices, e essa dimensão "
        "permanece fora do escopo desta avaliação.",
        "**Máquina única e não dedicada (validade de conclusão).** Toda a coleta ocorreu em um só "
        "*host*, compartilhado com o sistema operacional do usuário. Mitigou-se pela repetição "
        "das execuções sensíveis a ruído — o Lighthouse é reportado pela mediana de três "
        "execuções — e pela fixação dos parâmetros de cada ferramenta.",
        "**Representatividade dos cenários (validade de construção).** Os cenários derivam da "
        "árvore de utilidade do ATAM (Seção 6) e não de dados de uso real, inexistentes para uma "
        "plataforma ainda não operada em produção.",
        "**Avaliação conduzida pela equipe desenvolvedora (viés do avaliador).** Mitigou-se pela "
        "fixação dos critérios de aceitação **antes** da coleta (Seção 5.2) e pela regra de "
        "registro anti-fabricação (Seção 3.5), que obriga a reportar o resultado desfavorável com "
        "o mesmo rigor do favorável.",
    ])
    a.corpo(
        "Nem toda medição planejada é igualmente reproduzível por terceiros, e o Apêndice B "
        "(VIEIRA; COSTA, 2026b) classifica cada uma quanto a isso, explicitando as que **não** "
        "reproduzem e por quê — a cobertura do *baseline*, por depender de um estado de árvore "
        "anterior às correções, e as medições de tempo, que reproduzem a faixa e o veredito, não "
        "o valor exato."
    )


def secao_6(a):
    a.secao("6 Avaliação Arquitetural com ATAM")
    a.corpo(
        "Esta seção apresenta a aplicação do *Architecture Tradeoff Analysis Method* (ATAM) à "
        "arquitetura do Hubservi, conforme Clements, Kazman e Klein (2002). O ATAM complementa o "
        "plano de métricas (Seção 5): enquanto este coleta evidências quantitativas, o ATAM "
        "organiza a análise qualitativa das decisões arquiteturais, evidenciando riscos, pontos de "
        "sensibilidade e pontos de compromisso."
    )

    a.secao("6.1 Etapas do ATAM aplicadas ao estudo de caso")
    a.corpo(
        "O método foi conduzido nas etapas a seguir, adaptadas ao contexto de um estudo de caso "
        "conduzido pela própria equipe técnica:"
    )
    a.enumerada([
        "**Apresentação do ATAM** — alinhamento do método com os envolvidos.",
        "**Apresentação dos direcionadores de negócio** — intermediação de serviços, confiança "
        "entre partes e proteção de dados pessoais.",
        "**Apresentação da arquitetura** — conforme a Seção 4 (SPA + BaaS + Serverless).",
        "**Identificação das abordagens arquiteturais** — delegação de *backend* ao BaaS; "
        "autorização declarativa por RLS; regras de negócio em *triggers* e *views*; estado de "
        "servidor gerenciado por React Query no cliente.",
        "**Construção da árvore de utilidade** — Seção 6.2.",
        "**Análise das abordagens arquiteturais** — Seção 6.4.",
        "**Brainstorming e priorização de cenários** — refinamento dos cenários da árvore de "
        "utilidade.",
        "**Reanálise das abordagens** — à luz dos cenários priorizados e das métricas coletadas.",
        "**Apresentação dos resultados** — riscos, temas de risco e *tradeoffs* consolidados, "
        "reportados na Seção 6.4 com apoio das medições da Seção 7.",
    ])

    a.secao("6.2 Árvore de utilidade")
    a.corpo(
        "A árvore decompõe a utilidade geral em atributos de qualidade, refinamentos e cenários, "
        "cada qual rotulado por um par **(importância para o negócio, risco arquitetural)** em "
        "escala alto, médio e baixo — A, M e B (Figura 3)."
    )
    a.figura(
        "Árvore de utilidade dos cenários de atributo de qualidade",
        "docs/tcc/TCC-Hubservi.md, Seção 6.2 (diagrama Mermaid)",
    )
    a.corpo(
        "Os cenários priorizados são: **CS-1** isolamento de dados por RLS (A, A); **CS-2** "
        "anti-escalonamento de privilégio (A, M); **CS-3** não exposição de PII a anônimos (A, M); "
        "**CD-1** carregamento inicial da SPA (M, M); **CD-2** busca de serviços sob carga (A, M); "
        "**CM-1** baixa complexidade e duplicação (M, B); **CM-2** ausência de dependências "
        "circulares (M, B); **CC-1** máquina de estados do *booking* consistente (A, M); **CC-2** "
        "integridade referencial em cascata (A, B); e **CT-1** cobertura dos fluxos críticos "
        "(M, M)."
    )

    a.secao("6.3 Cenários de atributo de qualidade (formato de seis partes)")
    a.corpo(
        "Os cenários prioritários são especificados segundo a estrutura de Bass, Clements e Kazman "
        "(2012): fonte, estímulo, artefato, ambiente, resposta e medida."
    )
    a.corpo(
        "**CS-1 — Isolamento de dados por RLS.** *Fonte:* usuário autenticado mal-intencionado. "
        "*Estímulo:* requisição para ler ou alterar registros de outro usuário. *Artefato:* "
        "políticas RLS das tabelas `bookings`, `services` e `profiles`. *Ambiente:* operação "
        "normal. *Resposta:* a requisição é negada pelo banco. *Medida:* 0 acessos indevidos "
        "bem-sucedidos."
    )
    a.corpo(
        "**CS-2 — Anti-escalonamento de privilégio.** *Fonte:* usuário autenticado. *Estímulo:* "
        "tentativa de alterar o próprio `user_type`. *Artefato:* *trigger* "
        "`prevent_user_type_change()`. *Ambiente:* operação normal. *Resposta:* alteração "
        "rejeitada com exceção. *Medida:* 100% das tentativas bloqueadas."
    )
    a.corpo(
        "**CS-3 — Não exposição de PII.** *Fonte:* visitante não autenticado e usuário autenticado "
        "não relacionado. *Estímulo:* leitura de dados de perfil. *Artefato:* *view* "
        "`public_profiles` e RLS de `profiles`. *Ambiente:* operação normal. *Resposta:* apenas "
        "campos não sensíveis são retornados. *Medida:* 0 ocorrências de `email` ou `phone` "
        "expostos."
    )
    a.corpo(
        "**CD-2 — Busca de serviços sob carga.** *Fonte:* conjunto de usuários concorrentes. "
        "*Estímulo:* requisições simultâneas de busca e listagem. *Artefato:* API PostgREST e "
        "*view* `service_stats`. *Ambiente:* carga de pico planejada. *Resposta:* respostas "
        "corretas dentro do tempo-alvo. *Medida:* p95 menor ou igual a 800 ms; taxa de erro abaixo "
        "de 1%."
    )
    a.corpo(
        "**CC-1 — Máquina de estados do *booking* consistente.** *Fonte:* cliente ou prestador. "
        "*Estímulo:* tentativa de transição de status, válida ou inválida. *Artefato:* *trigger* "
        "`validate_booking_status_transition()`. *Ambiente:* operação normal. *Resposta:* "
        "transições válidas aplicadas e inválidas rejeitadas. *Medida:* 0 transições inválidas "
        "aceitas."
    )
    a.corpo(
        "Os demais cenários (CD-1, CM-1, CM-2, CC-2 e CT-1) seguem a mesma estrutura e estão "
        "associados às métricas da Seção 5."
    )

    a.secao("6.4 Pontos de sensibilidade, de compromisso e riscos")
    a.corpo(
        "A análise a seguir derivou da inspeção arquitetural (Seção 4) e foi confirmada pelas "
        "medições da Seção 7."
    )
    a.tabela(
        "Pontos de sensibilidade, pontos de compromisso, riscos e não riscos identificados",
        ["Tipo", "Descrição"],
        [
            ["**Ponto de sensibilidade**",
             "A correção da segurança é altamente sensível à completude e exatidão das políticas "
             "de RLS e dos *triggers*: uma política ausente ou mal especificada compromete CS-1 e "
             "CS-3 sem gerar erro funcional aparente. Os dois defeitos reportados na Seção 7.3.2 "
             "confirmam empiricamente essa sensibilidade."],
            ["**Ponto de sensibilidade**",
             "O desempenho de CD-2 é sensível à eficiência da *view* `service_stats` e à "
             "indexação das tabelas consultadas."],
            ["**Ponto de compromisso**",
             "A delegação total de autorização ao banco (RLS) favorece simplicidade e consistência "
             "(positivo para manutenibilidade e segurança), mas concentra o risco em um único "
             "mecanismo declarativo e desloca o esforço de teste para a fronteira cliente–banco "
             "(impacto em testabilidade)."],
            ["**Ponto de compromisso**",
             "A arquitetura SPA reduz acoplamento de *backend* e acelera o desenvolvimento, porém "
             "transfere carga de processamento e desempenho percebido para o cliente (tensão entre "
             "manutenibilidade e agilidade, de um lado, e desempenho — CD-1 — de outro)."],
            ["**Risco**",
             "O desempenho de carregamento inicial do *frontend* não atinge a meta (LCP de 3,17 s, "
             "acima de 2,5 s, mesmo após *code-splitting*), confirmando quantitativamente o ponto "
             "de compromisso CD-1 (Seção 7.3.5)."],
            ["**Não risco**",
             "A imutabilidade do `user_type` e a validação de `provider_id` por *trigger* fornecem "
             "defesa em profundidade contra escalonamento e *spoofing* de prestador, reduzindo o "
             "risco de CS-2."],
        ],
    )

    a.secao("6.5 Articulação com o plano de métricas")
    a.corpo(
        "Cada cenário da árvore de utilidade vincula-se a uma ou mais métricas da Seção 5, de modo "
        "que a coleta quantitativa forneceu a evidência para confirmar ou refutar a análise "
        "qualitativa do ATAM. Essa articulação entre ATAM (qualitativo) e ISO/IEC 25010 "
        "(quantitativo) constitui o procedimento de avaliação proposto como contribuição do "
        "trabalho."
    )


def secao_7(a):
    a.secao("7 Resultados")
    a.corpo(
        "Esta seção consolida os resultados do trabalho. Os objetivos específicos 1 a 4 "
        "(delimitação, modelagem, documentação e definição de cenários) foram concluídos e são "
        "reportados em 7.1 e 7.2. Os objetivos 5 a 9 (execução de testes automatizados, análise "
        "estática, testes de segurança e de desempenho, e análise) foram **executados** em julho "
        "de 2026: as medições são apresentadas em 7.3, e a síntese por atributo de qualidade em "
        "7.4."
    )
    a.corpo(
        "Todos os valores quantitativos aqui reportados derivam de execuções registradas de forma "
        "reprodutível, no ambiente caracterizado na Seção 5.4, onde cada medição remete a uma "
        "ferramenta e versão, ao ambiente, à data e a um arquivo de evidência (conforme a Seção "
        f"5.1). Nenhum número é apresentado sem evidência correspondente. O registro completo está "
        f"em {REGISTRO_URL} e as saídas brutas de cada execução, em {EVIDENCIAS_URL}; as capturas "
        "de tela dessas execuções compõem o Apêndice C (VIEIRA; COSTA, 2026c), ao qual cada "
        "subseção a seguir remete."
    )

    a.secao("7.1 Definição e delimitação do estudo")
    a.lista([
        "**Foco redefinido:** o trabalho foi reorientado da temática de recomendação para a "
        "**avaliação técnica da arquitetura de software**, mantendo o Hubservi como objeto de "
        "estudo. A recomendação permanece como módulo secundário (ordenação por popularidade e "
        "avaliação).",
        "**Objeto e instrumento (Seção 1.5):** o objeto de pesquisa é a avaliação; o Hubservi é o "
        "instrumento, desenvolvido pela equipe como pré-requisito metodológico — condição para o "
        "acesso irrestrito ao código, ao esquema e ao ambiente que a avaliação exige.",
        "**Modelo arquitetural consolidado:** **SPA + BaaS + Serverless**, corrigindo incoerência "
        "documental anterior que mencionava microsserviços (inexistentes no sistema real).",
        "**Atributos de qualidade selecionados** (ISO/IEC 25010): segurança, eficiência de "
        "desempenho, manutenibilidade (com ênfase em testabilidade) e confiabilidade.",
        "**Metodologia definida:** pesquisa aplicada, abordagem mista, objetivos "
        "exploratório-descritivos, estudo de caso com experimento técnico.",
    ])

    a.secao("7.2 Modelagem e documentação da arquitetura")
    a.corpo(
        "Foram produzidos e estão disponíveis no repositório do trabalho, em `docs/tcc/`, e "
        "reunidos no material suplementar (VIEIRA; COSTA, 2026a):"
    )
    a.lista([
        "descrição técnica completa da arquitetura, componentes, fluxos, persistência e regras de "
        "negócio (Seção 4);",
        "modelos UML — casos de uso, classes, sequência (autenticação, contratação e avaliação), "
        "componentes e implantação;",
        "modelo de processos de negócio (BPMN) para contratação e gerenciamento de *booking*;",
        "modelo de dados — DER e dicionário de dados, fiéis ao esquema real das *migrations*.",
    ])
    a.corpo(
        "Esses artefatos atendem aos objetivos específicos 2 e 3 e fundamentam a árvore de "
        "utilidade e os cenários do ATAM (Seção 6)."
    )

    a.secao("7.3 Resultados das medições")
    a.corpo(
        "A fase de execução configurou o ambiente reprodutível e coletou as métricas planejadas na "
        "Seção 5. O ambiente de teste de banco e de API foi levantado com uma instância "
        "**Supabase local** (contêiner Docker), contra a qual os cenários de autorização e de "
        "confiabilidade foram exercitados através da API real (PostgREST) — atendendo à exigência "
        "de teste de integração contra a API (Seção 5.2.1). Onde uma ferramenta nomeada no plano "
        "não estava disponível no ambiente, adotou-se um substituto da mesma classe, com a troca "
        "registrada: **k6 para autocannon** (teste de carga) e **Snyk para `npm audit`** (análise "
        "de composição de dependências)."
    )

    a.secao("7.3.1 Achado transversal de reprodutibilidade")
    a.corpo(
        "Ao reconstruir o banco a partir apenas das *migrations* versionadas, constatou-se que **o "
        "histórico não reproduzia um sistema funcional do zero**: faltavam os *grants* de "
        "privilégios de API aos papéis `anon`, `authenticated` e `service_role` no *schema* "
        "`public`, provocando `permission denied` antes mesmo da avaliação de RLS. O sistema "
        "funciona em produção porque esses privilégios foram aplicados fora do versionamento. "
        "Corrigiu-se com uma *migration* que torna o histórico auto-suficiente, sem alterar "
        "políticas de segurança. É um resultado característico do paradigma BaaS — parte da "
        "configuração vive no serviço gerenciado e escapa ao controle de versão — e só se revela "
        "sob uma avaliação conduzida em ambiente limpo e reprodutível. A Figura 4 mostra os "
        f"privilégios já concedidos após a correção (ver também Apêndice C, Figura {ref_c('P-09')})."
    )
    a.figura_arquivo(
        "Privilégios de API sobre `profiles` após a *migration* corretiva, verificados no banco "
        "reconstruído do zero",
        PRINTS / _CAPTURA["P-13"][1],
        fonte_captura("P-13"),
    )

    a.secao("7.3.2 Segurança")
    a.corpo(
        "Os cenários de autorização foram automatizados como testes de integração por papel. "
        "**Dois defeitos reais de autorização foram detectados, corrigidos e re-medidos**, "
        "seguindo o ciclo medir, detectar, corrigir e re-medir:"
    )
    a.lista([
        "**Exposição de PII a usuário autenticado:** a política de leitura de `profiles` (`USING "
        "(auth.uid() IS NOT NULL)`) permitia que qualquer usuário autenticado lesse `email` e "
        "`phone` de todos os perfis. Corrigido restringindo a leitura direta ao próprio registro; "
        "dados de contraparte seguem pela *view* `public_profiles`, sem PII.",
        "**Avaliação atribuída a prestador incorreto:** a política de inserção de *reviews* não "
        "validava que `provider_id` corresponde ao dono do serviço — assimetria em relação aos "
        "*bookings*, que já possuíam esse *trigger*. Corrigido com *trigger* de validação "
        "espelhado.",
    ])
    a.corpo(
        "As Figuras 5 e 6 registram o par antes/depois desses dois defeitos. Na primeira, os "
        "testes falham e a saída exibe o vazamento literal — o campo `email` de outro usuário "
        "devolvido pela API a um cliente que não deveria vê-lo; na segunda, os mesmos testes "
        "passam, contra o mesmo cenário, após as *migrations* corretivas. Preservar os dois "
        "estados é o que distingue um defeito detectado de uma suíte que nunca reprovou nada "
        "(Seção 3.5)."
    )
    a.figura_arquivo(
        "Antes da correção: falha dos testes de isolamento, com o vazamento de PII na saída",
        PRINTS / _CAPTURA["P-11"][1],
        fonte_captura("P-11"),
    )
    a.figura_arquivo(
        "Depois da correção: os mesmos testes aprovados, sem alteração no cenário",
        PRINTS / _CAPTURA["P-12"][1],
        fonte_captura("P-12"),
    )
    a.corpo(
        "Após as correções, a suíte de segurança (isolamento de perfis, serviços e *bookings*; "
        "bloqueio de escalonamento de `user_type`; regras de *review*) passa integralmente — **0 "
        "acessos indevidos nos cenários testados**, em 30 testes de integração distribuídos por "
        "nove arquivos. A análise de composição de dependências (`npm audit`) reportou 12 "
        "vulnerabilidades de severidade alta e 8 de severidade moderada, nenhuma crítica; destas, "
        "**apenas uma é de produção** (`react-router-dom`, *XSS via open redirect*, com correção "
        "disponível), sendo as demais ferramentas de *build* e de desenvolvimento, sem superfície "
        "de ataque em produção. O `supabase db lint` não acusou erros de *schema*. A varredura "
        "dinâmica (DAST/OWASP ZAP) permanece **pendente** de execução contra a URL de produção, "
        "por depender do ambiente de *hosting* real. As execuções correspondentes estão no "
        f"Apêndice C (Figuras {ref_c('P-10')}, {ref_c('P-17')} e {ref_c('P-18')})."
    )

    a.secao("7.3.3 Confiabilidade")
    a.corpo(
        "Verificados por testes de integração: a **máquina de estados do *booking*** rejeita "
        "transições inválidas (0 transições inválidas aceitas); a **integridade referencial** em "
        "exclusão de serviço não deixa registros órfãos (cascata); e o **fluxo crítico ponta a "
        "ponta** (autenticação, serviço, *booking* e avaliação) completa com sucesso no caso "
        f"válido. Atende aos critérios da Seção 5.2.5 (ver Apêndice C, Figura {ref_c('P-10')})."
    )

    a.secao("7.3.4 Testabilidade")
    a.corpo(
        "A suíte automatizada foi ampliada de 11 para **44 testes unitários e de componente**, "
        "além de 30 testes de integração. A **cobertura de linhas subiu de 18,0% para 32,0%** no "
        "agregado, concentrada nos módulos críticos, que passaram a **82% a 100%** (contexto de "
        "autenticação, formulários, camada de acesso a dados e *schemas* de validação). Definiu-se "
        "o limiar antes em aberto (Seção 5.2.3): piso global anti-regressão de 30% e piso de 75% "
        "nos módulos críticos, enumerados nominalmente e verificados automaticamente. "
        "Introduziu-se uma *factory* de *mock* compartilhada, reduzindo o esforço de escrever "
        "novos testes — evidência de que a testabilidade da arquitetura, atributo sob avaliação, "
        "melhorou."
    )
    a.figura_arquivo(
        "Cobertura de testes ao final da ampliação da suíte",
        PRINTS / _CAPTURA["P-03"][1],
        fonte_captura("P-03"),
    )
    a.corpo(
        "O estado inicial não dispõe de captura equivalente, e a assimetria é registrada em vez "
        "de contornada: a cobertura de 18,03% foi medida sobre uma árvore de trabalho anterior às "
        "correções e é, conforme o Apêndice B (VIEIRA; COSTA, 2026b), a única medição do trabalho "
        "que **não se reproduz em *commit* algum**. O que se preserva do estado inicial é a suíte "
        f"do *baseline*, com seus 11 testes (Apêndice C, Figura {ref_c('P-02')}), não o "
        "percentual."
    )

    a.secao("7.3.5 Eficiência de desempenho")
    a.corpo("Distinguem-se duas camadas:")
    a.lista([
        "**Backend sob carga (atende):** a listagem de serviços respondeu, sob 30 conexões "
        "concorrentes por 20 s, com **latência de cauda (p97,5) de 253 ms** (critério de 800 ms "
        "em p95, portanto atendido com folga) e **0% de "
        "erro** em 44.413 requisições, a uma vazão aproximada de 2.221 requisições por segundo.",
        "**Frontend e carregamento inicial (não atende):** o Lighthouse (mediana de três "
        "execuções, perfil móvel com limitação de CPU e de rede) registrou *performance score* de "
        "85 e **LCP de 3,49 s** (critérios de 90 e de 2,5 s). Diagnosticou-se *bundle* único de "
        "679 KB com rotas carregadas de forma *eager*; aplicou-se *code-splitting* por rota, "
        "reduzindo o *bundle* inicial em 28% (para 489 KB) e melhorando os índices (score de 88 e "
        "LCP de 3,17 s), que **permanecem abaixo da meta**. O TBT caiu de 11 ms para 0 ms e o CLS "
        "manteve-se em 0,000.",
    ])
    a.corpo(
        "As Figuras 8 e 9 apresentam os relatórios do Lighthouse antes e depois do "
        "*code-splitting*, na execução mediana de cada rodada de três. A comparação torna visível "
        "tanto o ganho quanto o seu limite: os indicadores melhoram na direção esperada, e ainda "
        "assim o LCP permanece acima do critério de 2,5 s."
    )
    a.figura_arquivo(
        "Carregamento inicial antes do *code-splitting*: *score* de 85 e LCP de 3,49 s",
        PRINTS / _CAPTURA["P-14"][1],
        fonte_captura("P-14"),
    )
    a.figura_arquivo(
        "Carregamento inicial depois do *code-splitting*: *score* de 88 e LCP de 3,17 s",
        PRINTS / _CAPTURA["P-15"][1],
        fonte_captura("P-15"),
    )
    a.corpo(
        "O contraste entre a API rápida e o carregamento lento localiza o gargalo de desempenho no "
        "*frontend* (peso do *bundle*), e não no *backend*. O ensaio de carga que sustenta o "
        f"resultado de *backend* consta do Apêndice C, Figura {ref_c('P-16')}; note-se que ele "
        "corre sobre o volume de dados declarado na Seção 5.4, e vale como piso, não como "
        "projeção de operação real (Seção 5.5)."
    )

    a.secao("7.3.6 Manutenibilidade")
    a.corpo(
        "A **modularização é sólida**: **0 dependências circulares**, confirmado por duas "
        "ferramentas independentes (Madge, sobre 85 arquivos, e dependency-cruiser, sobre 93 "
        "módulos), com grafo de acoplamento saudável (núcleo estável e folhas voláteis). A "
        "**higiene de código está abaixo do ideal**: o *lint* não está zerado (19 violações na "
        "configuração atual e 25 sob configuração recomendada, incluindo funções com complexidade "
        "ciclomática elevada), e a **duplicação na camada de interface é de 4,55%** (acima da meta "
        "de 3%), com causa localizada — os dois painéis, de cliente e de prestador, compartilham "
        f"61 linhas quase idênticas. As execuções constam do Apêndice C (Figuras {ref_c('P-04')}, "
        f"{ref_c('P-05')}, {ref_c('P-07')} e {ref_c('P-08')})."
    )

    a.secao("7.4 Síntese dos resultados")
    a.corpo(
        "A avaliação técnica foi conduzida de ponta a ponta sobre um ambiente reprodutível e emite "
        "vereditos **por atributo e por camada**, em vez de um juízo único."
    )
    a.tabela(
        "Síntese dos resultados por atributo de qualidade da ISO/IEC 25010",
        ["Atributo (ISO/IEC 25010)", "Resultado"],
        [
            ["Segurança",
             "Autorização declarativa (RLS e *triggers*) **atende** após a correção de dois "
             "defeitos reais; 1 CVE de produção a tratar; DAST pendente contra produção"],
            ["Confiabilidade",
             "**Atende** (máquina de estados, integridade referencial e fluxo crítico)"],
            ["Testabilidade",
             "**Evoluiu** (cobertura de 18% para 32%; módulos críticos entre 82% e 100%, com "
             "limiar protegido)"],
            ["Eficiência de desempenho",
             "*Backend* **atende** (p95 de 253 ms); *frontend* **não atende** (LCP de 3,17 s após "
             "otimização)"],
            ["Manutenibilidade",
             "Estrutura **atende** (0 ciclos); higiene de código **abaixo do ideal** (*lint* e "
             "duplicação nos painéis)"],
        ],
    )
    a.corpo(
        "O trabalho responde à questão de pesquisa demonstrando um **procedimento reprodutível de "
        "avaliação técnica** aplicável a arquiteturas BaaS/Serverless: modelagem fiel, cenários "
        "derivados dos atributos de qualidade, medição instrumentada com registro de evidências e "
        "o ciclo de detecção e correção de defeitos. O resultado não é um atestado de que o "
        "sistema é uniformemente bom — três defeitos reais foram encontrados e o desempenho de "
        "*frontend* não atinge a meta —, e é justamente essa capacidade de **localizar "
        "deficiências com precisão** que evidencia o valor do método. A reprodutibilidade é "
        "assegurada pelo registro em `docs/tcc/medicoes/` e pela integração contínua, que executa "
        "os *gates* automatizados a cada alteração."
    )


# A Seção "8 Cronograma" foi removida nesta versão final. Um cronograma de etapas
# futuras é elemento de projeto de pesquisa, não de relato conclusivo: o trabalho
# terminou, e a tabela de meses com situação "Planejado" contradizia a Seção 7, que
# já reporta as medições executadas. As pendências que ali constavam (DAST contra a
# URL de produção) passaram a ser tratadas como limitação na conclusão, que é o lugar
# metodologicamente correto para elas. A conclusão, antes Seção 9, passa a ser a 8.


def secao_8(a):
    a.secao("8 Conclusão")
    a.corpo(
        "Este trabalho investigou como avaliar tecnicamente uma arquitetura web baseada em BaaS e "
        "*Serverless* por meio de métricas e testes de Engenharia de Software. A resposta oferecida "
        "à questão de pesquisa é um **procedimento de avaliação arquitetural** que articula três "
        "elementos: o modelo de qualidade da ISO/IEC 25010 (2011), que fornece a taxonomia dos "
        "atributos; o método ATAM (CLEMENTS; KAZMAN; KLEIN, 2002), que organiza a análise "
        "qualitativa em cenários, riscos e pontos de compromisso; e um conjunto instrumentado de "
        "ferramentas de teste automatizado, análise estática, desempenho e segurança, que produz "
        "a evidência quantitativa. O procedimento foi aplicado integralmente à plataforma "
        "Hubservi, adotada como instrumento por permitir o acesso irrestrito ao código, ao esquema "
        "e ao ambiente que a avaliação exige."
    )
    a.corpo(
        "Os resultados demonstram que o procedimento é capaz de **localizar deficiências com "
        "precisão**, e não apenas de emitir um juízo agregado. Três defeitos reais foram "
        "detectados, corrigidos e re-medidos: a exposição de dados pessoais a qualquer usuário "
        "autenticado, decorrente de política de RLS excessivamente permissiva; a ausência de "
        "validação do prestador na inserção de avaliações; e a incapacidade de o histórico de "
        "*migrations* reconstruir um sistema funcional do zero, por privilégios aplicados fora do "
        "controle de versão. Os três são achados característicos do paradigma BaaS — todos "
        "residem em configuração declarativa ou no serviço gerenciado, e nenhum se manifestava "
        "como erro funcional aparente na aplicação."
    )
    a.corpo(
        "A avaliação também evidenciou que vereditos por camada são mais informativos do que um "
        "veredito único: enquanto o *backend* atende folgadamente ao critério de desempenho, o "
        "carregamento inicial do *frontend* permanece abaixo da meta mesmo após otimização por "
        "*code-splitting*, o que confirma quantitativamente o ponto de compromisso identificado na "
        "análise ATAM — a arquitetura SPA transfere ao cliente parte do desempenho percebido. De "
        "modo análogo, a manutenibilidade apresenta estrutura sólida, sem dependências circulares, "
        "convivendo com higiene de código abaixo do ideal em *lint* e duplicação."
    )
    a.corpo(
        "A principal contribuição é metodológica: um roteiro reprodutível, com registro de "
        "evidências e critérios fixados antes da coleta, transferível a outras aplicações que "
        "adotem o mesmo paradigma. A contribuição prática, para a plataforma avaliada, materializa-se "
        "nas correções aplicadas e nos *gates* automatizados incorporados à integração contínua. "
        "Como **limitações**, registram-se a varredura dinâmica de segurança (DAST) ainda pendente "
        "de execução contra a URL de produção, a condução da avaliação pela própria equipe "
        "desenvolvedora — o que exigiu explicitar critérios previamente para mitigar viés — e a "
        "generalização restrita a um único caso. Como **trabalhos futuros**, propõem-se a execução "
        "do DAST em ambiente publicado, a aplicação do mesmo procedimento a outras plataformas "
        "BaaS para avaliar sua transferibilidade, a redução da duplicação identificada entre os "
        "painéis e a continuidade da otimização de carregamento do *frontend* até o alcance da "
        "meta de LCP."
    )


def referencias(a):
    # Apenas as obras efetivamente citadas no corpo. As quatro entradas da lista
    # "complementares sugeridas (opcional)" do consolidado — Fowler, ISO/IEC 25023,
    # Supabase e PostgreSQL — foram removidas por não terem citação correspondente.
    a.secao("Referências", primaria=True)
    par = a.doc.add_paragraph()
    par.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    par.paragraph_format.first_line_indent = RECUO
    entradas = [
        ("ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. ", "NBR 6023",
         ": informação e documentação: referências: elaboração. Rio de Janeiro: ABNT, 2018."),
        ("ASSOCIAÇÃO BRASILEIRA DE NORMAS TÉCNICAS. ", "NBR 10520",
         ": informação e documentação: citações em documentos: apresentação. Rio de Janeiro: "
         "ABNT, 2023."),
        ("BASS, Len; CLEMENTS, Paul; KAZMAN, Rick. ", "Software architecture in practice",
         ". 3rd ed. Upper Saddle River: Addison-Wesley, 2012."),
        ("CLEMENTS, Paul; KAZMAN, Rick; KLEIN, Mark. ", "Evaluating software architectures",
         ": methods and case studies. Boston: Addison-Wesley, 2002."),
        ("INTERNATIONAL ORGANIZATION FOR STANDARDIZATION; INTERNATIONAL ELECTROTECHNICAL "
         "COMMISSION. ", "ISO/IEC 25010",
         ": systems and software engineering — Systems and software Quality Requirements and "
         "Evaluation (SQuaRE) — System and software quality models. Geneva: ISO, 2011."),
        ("PRESSMAN, Roger S.; MAXIM, Bruce R. ", "Engenharia de software",
         ": uma abordagem profissional. 8. ed. Porto Alegre: AMGH, 2016."),
        ("SOMMERVILLE, Ian. ", "Engenharia de software",
         ". 9. ed. São Paulo: Pearson Prentice Hall, 2011."),
        # Material suplementar dos próprios autores: os diagramas e o protocolo de
        # reprodução saíram do corpo do artigo e passaram a ser documentos citáveis.
        ("VIEIRA, Pedro Conrado Fernandes; COSTA, Richardy Gabriel Rodrigues da. ",
         "Apêndice A — Diagramas (UML, BPMN e DER) e dicionário de dados",
         f": material suplementar. Franca: Uni-FACEF, 2026a. Disponível em: {APENDICE_A_URL}. "
         "Acesso em: 20 ago. 2026."),
        ("VIEIRA, Pedro Conrado Fernandes; COSTA, Richardy Gabriel Rodrigues da. ",
         "Apêndice B — Reprodução das medições",
         f": material suplementar. Franca: Uni-FACEF, 2026b. Disponível em: {APENDICE_B_URL}. "
         "Acesso em: 20 ago. 2026."),
        ("VIEIRA, Pedro Conrado Fernandes; COSTA, Richardy Gabriel Rodrigues da. ",
         "Apêndice C — Evidências de execução",
         f": material suplementar. Franca: Uni-FACEF, 2026c. Disponível em: {APENDICE_C_URL}. "
         "Acesso em: 28 ago. 2026."),
    ]
    for i, (antes, titulo, depois) in enumerate(entradas):
        if i:
            par.add_run().add_break()
        a._run(par, antes)
        a._run(par, titulo, bold=True)
        a._run(par, depois)


def apendice(a):
    """Conteúdo do Apêndice A — gerado como DOCUMENTO PRÓPRIO, não como seção do artigo.

    O artigo não leva apêndice: os diagramas são artefatos do instrumento (a plataforma),
    não resultados da avaliação (o objeto de pesquisa). Publicá-los à parte preserva essa
    distinção e mantém o corpo do texto no tamanho de um artigo.

    A numeração usa o prefixo "A." e é autocontida, de modo que inserir ou remover figuras
    no corpo do artigo não desloque as referências cruzadas daqui.
    """
    a.prefixo = "A."
    # As imagens vêm da mesma origem que as do artigo, depois das três do corpo
    # (Figuras 1 a 3). Sem pular, o apêndice começaria pela imagem do corpo e trocaria
    # todas as legendas em silêncio.
    a.pular_imagens(3)

    a.titulo_artigo("Apêndice A — Diagramas (UML, BPMN e DER) e dicionário de dados")
    a.centro("MATERIAL SUPLEMENTAR")
    a.vazio()
    a.corpo(
        "Material suplementar do artigo *Avaliação Técnica de uma Arquitetura Web baseada em "
        "BaaS/Serverless para Intermediação de Serviços: um estudo de caso da plataforma "
        "Hubservi*, de Pedro Conrado Fernandes Vieira e Richardy Gabriel Rodrigues da Costa, "
        "sob orientação do Prof. Daniel Facciolo Pires (Uni-FACEF, 2026)."
    )
    a.corpo(
        "Os diagramas a seguir foram recuperados do código-fonte e das *migrations*; o código de "
        f"diagramação que os origina está versionado em `docs/tcc/diagramas/`, no repositório do "
        f"trabalho ({REPO_URL}). A versão navegável deste documento, com os diagramas renderizados "
        f"pelo próprio GitHub, está em {APENDICE_A_URL}."
    )
    a.corpo(
        "O protocolo de reprodução das medições reportadas na Seção 7 do artigo — incluindo o que "
        "não é reproduzível e por quê — está no Apêndice B (VIEIRA; COSTA, 2026b), disponível em "
        f"{APENDICE_B_URL}."
    )

    diagramas = [
        ("Diagrama Entidade-Relacionamento (DER)", "docs/tcc/diagramas/der.md"),
        ("Diagrama de Casos de Uso", "docs/tcc/diagramas/caso-de-uso.md"),
        ("Diagrama de Classes", "docs/tcc/diagramas/classes.md"),
        ("Diagrama de Sequência — Autenticação",
         "docs/tcc/diagramas/sequencia-autenticacao.md"),
        ("Diagrama de Sequência — Contratação (booking)",
         "docs/tcc/diagramas/sequencia-contratacao.md"),
        ("Diagrama de Sequência — Avaliação (review)",
         "docs/tcc/diagramas/sequencia-avaliacao.md"),
        ("Diagrama de Componentes", "docs/tcc/diagramas/componentes.md"),
        ("Diagrama de Implantação", "docs/tcc/diagramas/implantacao.md"),
        ("BPMN — Contratação de serviço", "docs/tcc/diagramas/bpmn-contratacao.md"),
        ("BPMN e máquina de estados — Gerenciamento de booking",
         "docs/tcc/diagramas/bpmn-gerenciamento-booking.md"),
    ]
    for legenda, origem in diagramas:
        a.figura(legenda, origem)

    a.secao("A.1 Dicionário de dados")
    a.corpo(
        "Derivado das *migrations* em `supabase/migrations/` (*migration* inicial "
        "`20260303232457` e posteriores). Tipos conforme PostgreSQL."
    )
    a.tabela(
        "Tipos enumerados do esquema",
        ["Tipo", "Valores"],
        [
            ["`user_type`", "`client`, `provider`"],
            ["`price_type`", "`fixed`, `hourly`, `negotiable`"],
            ["`booking_status`", "`pending`, `accepted`, `completed`, `rejected`, `cancelled`"],
        ],
    )
    a.tabela(
        "Tabela `profiles`",
        ["Coluna", "Tipo", "Restrições", "Descrição"],
        [
            ["`id`", "uuid", "PK; FK para `auth.users(id)` ON DELETE CASCADE",
             "Identificador do usuário"],
            ["`email`", "text", "UNIQUE, NOT NULL", "E-mail (sincronizado do Auth) — **PII**"],
            ["`full_name`", "text", "NOT NULL, DEFAULT ''", "Nome de exibição"],
            ["`phone`", "text", "DEFAULT ''", "Telefone — **PII**"],
            ["`avatar_url`", "text", "DEFAULT ''", "URL do avatar"],
            ["`user_type`", "user_type", "NOT NULL, DEFAULT 'client'; imutável (*trigger*)",
             "Papel do usuário"],
            ["`created_at`", "timestamptz", "NOT NULL, DEFAULT now()", "Criação"],
            ["`updated_at`", "timestamptz", "NOT NULL, DEFAULT now()", "Atualização (*trigger*)"],
        ],
    )
    a.tabela(
        "Tabela `categories`",
        ["Coluna", "Tipo", "Restrições", "Descrição"],
        [
            ["`id`", "uuid", "PK, DEFAULT gen_random_uuid()", "Identificador"],
            ["`name`", "text", "UNIQUE, NOT NULL", "Nome da categoria"],
            ["`description`", "text", "DEFAULT ''", "Descrição"],
            ["`icon`", "text", "DEFAULT ''", "Ícone (UI)"],
            ["`created_at`", "timestamptz", "NOT NULL, DEFAULT now()", "Criação"],
        ],
    )
    a.tabela(
        "Tabela `services`",
        ["Coluna", "Tipo", "Restrições", "Descrição"],
        [
            ["`id`", "uuid", "PK, DEFAULT gen_random_uuid()", "Identificador"],
            ["`provider_id`", "uuid", "NOT NULL, FK para `profiles(id)` ON DELETE CASCADE",
             "Prestador dono"],
            ["`category_id`", "uuid", "NOT NULL, FK para `categories(id)` ON DELETE RESTRICT",
             "Categoria"],
            ["`title`", "text", "NOT NULL", "Título"],
            ["`description`", "text", "NOT NULL, DEFAULT ''", "Descrição"],
            ["`price_min`", "numeric(10,2)", "NOT NULL, DEFAULT 0", "Preço mínimo"],
            ["`price_max`", "numeric(10,2)",
             "CHECK (`price_max` IS NULL OR `price_max` maior ou igual a `price_min`)",
             "Preço máximo (opcional)"],
            ["`price_type`", "price_type", "NOT NULL, DEFAULT 'fixed'", "Modelo de preço"],
            ["`location`", "text", "DEFAULT ''", "Localidade"],
            ["`is_active`", "boolean", "NOT NULL, DEFAULT true", "Visibilidade pública"],
            ["`created_at`", "timestamptz", "NOT NULL, DEFAULT now()", "Criação"],
            ["`updated_at`", "timestamptz", "NOT NULL, DEFAULT now()", "Atualização (*trigger*)"],
        ],
    )
    a.corpo(
        "Índices de `services`: `provider_id`, `category_id`, `is_active` e índice GIN de busca "
        "textual sobre `title`."
    )
    a.tabela(
        "Tabela `bookings`",
        ["Coluna", "Tipo", "Restrições", "Descrição"],
        [
            ["`id`", "uuid", "PK, DEFAULT gen_random_uuid()", "Identificador"],
            ["`service_id`", "uuid", "NOT NULL, FK para `services(id)` ON DELETE CASCADE",
             "Serviço"],
            ["`client_id`", "uuid", "NOT NULL, FK para `profiles(id)` ON DELETE CASCADE",
             "Cliente"],
            ["`provider_id`", "uuid", "NOT NULL, FK para `profiles(id)` ON DELETE CASCADE",
             "Prestador (deve coincidir com o dono do serviço — *trigger*)"],
            ["`status`", "booking_status",
             "NOT NULL, DEFAULT 'pending'; transições validadas (*trigger*)", "Estado"],
            ["`message`", "text", "DEFAULT ''", "Mensagem da solicitação"],
            ["`scheduled_date`", "timestamptz", "NULL", "Data agendada"],
            ["`created_at`", "timestamptz", "NOT NULL, DEFAULT now()", "Criação"],
            ["`updated_at`", "timestamptz", "NOT NULL, DEFAULT now()", "Atualização (*trigger*)"],
        ],
    )
    a.corpo("Índices de `bookings`: `client_id`, `provider_id` e `service_id`.")
    a.tabela(
        "Tabela `reviews`",
        ["Coluna", "Tipo", "Restrições", "Descrição"],
        [
            ["`id`", "uuid", "PK, DEFAULT gen_random_uuid()", "Identificador"],
            ["`service_id`", "uuid", "NOT NULL, FK para `services(id)` ON DELETE CASCADE",
             "Serviço avaliado"],
            ["`client_id`", "uuid", "NOT NULL, FK para `profiles(id)` ON DELETE CASCADE",
             "Autor (cliente)"],
            ["`provider_id`", "uuid", "NOT NULL, FK para `profiles(id)` ON DELETE CASCADE",
             "Prestador avaliado (validado por *trigger*)"],
            ["`rating`", "integer", "NOT NULL, CHECK (1 a 5)", "Nota"],
            ["`comment`", "text", "NOT NULL, DEFAULT ''", "Comentário"],
            ["`created_at`", "timestamptz", "NOT NULL, DEFAULT now()", "Criação"],
            ["—", "—", "UNIQUE (`service_id`, `client_id`)",
             "Uma avaliação por cliente por serviço"],
        ],
    )
    a.corpo("Índices de `reviews`: `service_id` e `client_id`.")
    a.tabela(
        "Views do esquema",
        ["View", "Colunas", "Descrição"],
        [
            ["`service_stats`",
             "`service_id`, `review_count` (int), `average_rating` (numeric(3,2))",
             "Agregação de avaliações por serviço; `security_invoker = true`"],
            ["`public_profiles`", "`id`, `full_name`, `avatar_url`, `user_type`, `created_at`",
             "Projeção sem PII (omite `email` e `phone`) para consumo anônimo"],
        ],
    )


def apendice_c(a):
    """Conteúdo do Apêndice C — Evidências de execução, como DOCUMENTO PRÓPRIO.

    Mesma lógica dos Apêndices A e B: material suplementar citável, com numeração
    autocontida (prefixo "C."), de modo que inserir uma figura no corpo do artigo não
    desloque nenhuma referência daqui.

    A fronteira com o Apêndice B é deliberada e vale dizê-la: o B ensina **como
    reproduzir** cada medição; o C mostra **o que se viu** quando ela foi executada.
    Um é procedimento, o outro é registro — não se substituem.
    """
    a.prefixo = "C."
    # Nenhuma imagem vem do .docx de origem: todas as capturas são lidas do repositório.

    a.titulo_artigo("Apêndice C — Evidências de execução")
    a.centro("MATERIAL SUPLEMENTAR")
    a.vazio()
    a.corpo(
        "Material suplementar do artigo *Avaliação Técnica de uma Arquitetura Web baseada em "
        "BaaS/Serverless para Intermediação de Serviços: um estudo de caso da plataforma "
        "Hubservi*, de Pedro Conrado Fernandes Vieira e Richardy Gabriel Rodrigues da Costa, "
        "sob orientação do Prof. Daniel Facciolo Pires (Uni-FACEF, 2026)."
    )
    a.corpo(
        "Este apêndice reúne as capturas de tela de todas as execuções de ferramenta que "
        "sustentam os resultados da Seção 7 do artigo. As mais relevantes — as que evidenciam os "
        "três defeitos detectados e os pares antes/depois de cobertura e de desempenho — são "
        "reproduzidas também no corpo do artigo; o conjunto completo fica aqui, para que a "
        "verificação de qualquer número reportado não dependa de acesso ao ambiente de execução."
    )
    a.corpo(
        "As saídas brutas correspondentes, citadas na linha de fonte de cada figura, estão "
        f"versionadas em {EVIDENCIAS_URL}, e o registro de cada medição — com ferramenta, versão, "
        f"critério, valor e veredito — em {REGISTRO_URL}. O procedimento para reexecutar cada uma "
        "delas é o do Apêndice B (VIEIRA; COSTA, 2026b)."
    )
    a.corpo(
        "**Nota de método, válida para todas as figuras.** As capturas **reencenam na tela as "
        "saídas preservadas** das execuções de 15 e 16 de julho de 2026, no *commit* `ad89e6c`. "
        "Nenhuma medição foi reexecutada para produzi-las, e é isso que garante que cada número "
        "visível nas imagens seja idêntico ao reportado no artigo. A reencenação é literal: as "
        "saídas de terminal são impressas byte a byte, com as sequências de cor originais, e os "
        "relatórios do ESLint e do Lighthouse são renderizados pelo relatório oficial da própria "
        "ferramenta, a partir dos arquivos JSON arquivados. Reexecutar as ferramentas hoje "
        "produziria números diferentes dos do artigo — e a saída que evidencia a exposição de PII "
        f"(Figura {ref_c('P-11')}) sequer é reproduzível, já que o defeito foi corrigido."
    )
    a.corpo(
        "**Uma lacuna declarada.** O catálogo de capturas do repositório prevê 18 telas; 17 estão "
        "reproduzidas aqui. Falta a do Madge (medição M-03, ausência de dependências circulares), "
        "que não chegou a ser capturada. A medição não fica sem evidência: ela é atestada pelo "
        "arquivo bruto `evidencias/2026-07-15/madge-circular.txt` e confirmada de forma "
        f"independente por segunda ferramenta, o dependency-cruiser (Figura {ref_c('P-08')})."
    )

    a.tabela(
        "Índice das evidências de execução",
        ["Figura", "Captura", "Ferramenta", "Saída bruta preservada"],
        [[f"C.{i}", pid, ferramenta, f"`evidencias/{arquivo}`"]
         for i, (pid, _, _, ferramenta, _, arquivo) in enumerate(CAPTURAS, 1)],
        fonte="Fonte: elaborado pelos autores (2026).",
    )

    for pid, arquivo, legenda, _, _, _ in CAPTURAS:
        a.figura_arquivo(legenda, PRINTS / arquivo, fonte_captura(pid))


def gerar_apendice_c_md():
    """Escreve a versão navegável do Apêndice C, em Markdown, no repositório.

    É o arquivo para onde aponta a URL citada no artigo (APENDICE_C_URL): sem ele, a
    referência VIEIRA; COSTA (2026c) seria um link morto. Gerar os dois formatos da
    mesma lista `CAPTURAS` garante que a numeração C.x do .docx e a do .md nunca
    divirjam — que é o defeito clássico de material suplementar mantido à mão.
    """
    destino = RAIZ / "docs" / "tcc" / "apendice-c-evidencias.md"
    L = []
    L.append("# Apêndice C — Evidências de execução\n")
    L.append(
        "**Material suplementar** do artigo *Avaliação Técnica de uma Arquitetura Web baseada "
        "em BaaS/Serverless para Intermediação de Serviços: um estudo de caso da plataforma "
        "Hubservi*.\n"
    )
    L.append("Pedro Conrado Fernandes Vieira · Richardy Gabriel Rodrigues da Costa  ")
    L.append("Graduandos em Engenharia de Software — Uni-FACEF  ")
    L.append("Orientador: Prof. Daniel Facciolo Pires\n")
    L.append("> **Arquivo gerado** por `docs/tcc/gerar-artigo-docx.py`. Para alterar legendas ou "
             "a ordem das figuras, edite a lista `CAPTURAS` naquele script e regenere — assim o "
             "Markdown e o `.docx` continuam com a mesma numeração.\n")
    L.append("---\n")
    L.append("## O que este documento é\n")
    L.append(
        "As capturas de tela de **todas** as execuções de ferramenta que sustentam a Seção 7 do "
        "artigo. As mais relevantes — os três defeitos detectados e os pares antes/depois de "
        "cobertura e de desempenho — aparecem também no corpo do artigo; o conjunto completo "
        "fica aqui, para que verificar qualquer número reportado não dependa de acesso ao "
        "ambiente de execução.\n"
    )
    L.append(
        "A fronteira com o [Apêndice B](apendice-b-reproducao.md) é deliberada: o **B** ensina "
        "*como reproduzir* cada medição; o **C** mostra *o que se viu* quando ela foi executada. "
        "Um é procedimento, o outro é registro.\n"
    )
    L.append("---\n")
    L.append("## Nota de método — vale como legenda de todas as figuras\n")
    L.append(
        "As imagens **reencenam na tela as saídas preservadas** das execuções de **15 e 16 de "
        "julho de 2026**, no commit `ad89e6c`. **Nenhuma medição foi reexecutada para "
        "produzi-las** — é isso que garante que cada número visível seja idêntico ao reportado "
        "no artigo. A reencenação é literal: as saídas de terminal são impressas byte a byte, "
        "com as sequências de cor originais, e os relatórios do ESLint e do Lighthouse são "
        "renderizados pelo relatório oficial da própria ferramenta a partir dos JSON "
        "arquivados.\n"
    )
    L.append(
        "Reexecutar as ferramentas hoje produziria números diferentes dos do artigo — e a saída "
        f"que evidencia a exposição de PII (Figura {ref_c('P-11')}) sequer é reproduzível, já "
        "que o defeito foi corrigido. O procedimento de captura está em "
        "[`medicoes/evidencias/prints/README.md`](medicoes/evidencias/prints/README.md).\n"
    )
    L.append("### Uma lacuna declarada\n")
    L.append(
        "O catálogo do repositório prevê 18 telas; **17** estão aqui. Falta a do Madge (medição "
        "M-03, ausência de dependências circulares), que não chegou a ser capturada. A medição "
        "não fica sem evidência: é atestada pelo arquivo bruto "
        "[`madge-circular.txt`](medicoes/evidencias/2026-07-15/madge-circular.txt) e confirmada "
        f"de forma independente pelo dependency-cruiser (Figura {ref_c('P-08')}).\n"
    )
    L.append("---\n")
    L.append("## Índice\n")
    L.append("| Figura | Captura | Ferramenta | Saída bruta preservada |")
    L.append("|---|---|---|---|")
    for i, (pid, _, _, ferramenta, _, arquivo) in enumerate(CAPTURAS, 1):
        L.append(f"| C.{i} | {pid} | {ferramenta} | "
                 f"[`{arquivo}`](medicoes/evidencias/{arquivo}) |")
    L.append("")
    L.append("---\n")
    L.append("## Figuras\n")
    for i, (pid, arquivo, legenda, ferramenta, data, bruto) in enumerate(CAPTURAS, 1):
        L.append(f"### Figura C.{i} — {legenda}\n")
        L.append(f"![Figura C.{i}](medicoes/evidencias/prints/{arquivo})\n")
        L.append(f"*Fonte: {ferramenta}; execução de {data}; saída bruta preservada em "
                 f"[`{bruto}`](medicoes/evidencias/{bruto}).*\n")
    L.append("---\n")
    L.append("## Material suplementar relacionado\n")
    L.append("| Documento | Conteúdo |")
    L.append("|---|---|")
    L.append("| [Apêndice A — Diagramas e dicionário de dados](apendice-a-diagramas.md) | "
             "UML, BPMN, DER e dicionário de dados |")
    L.append("| [Apêndice B — Reprodução das medições](apendice-b-reproducao.md) | "
             "Como reexecutar cada medição, e o que não reproduz |")
    L.append("| [Registro de medições](medicoes/registro-medicoes.md) | "
             "Tabela mestra: valor, ferramenta, versão, evidência, veredito |")
    L.append("| [Evidências brutas](medicoes/evidencias/) | "
             "Saídas originais das ferramentas, por data |")
    L.append("")
    L.append("### Como citar\n")
    L.append("> VIEIRA, Pedro Conrado Fernandes; COSTA, Richardy Gabriel Rodrigues da. "
             "**Apêndice C — Evidências de execução**: material suplementar. Franca: "
             f"Uni-FACEF, 2026. Disponível em: {APENDICE_C_URL}. Acesso em: [data].")

    destino.write_text("\n".join(L) + "\n", encoding="utf-8")
    return destino


def main():
    # 1. O artigo — sem apêndice, terminando nas Referências.
    a = Artigo(ORIGEM)
    capa(a)
    resumo(a)
    secao_1(a)
    secao_2(a)
    secao_3(a)
    secao_4(a)
    secao_5(a)
    secao_6(a)
    secao_7(a)
    secao_8(a)
    referencias(a)
    a.salvar(SAIDA)
    print(f"Gerado: {SAIDA}")
    print(f"  parágrafos: {len(a.doc.paragraphs)} | tabelas: {a.n_tabela} | figuras: {a.n_figura}")

    # 2. O Apêndice A — documento próprio, com numeração A.x e as imagens 4 a 13.
    b = Artigo(ORIGEM)
    apendice(b)
    b.salvar(SAIDA_APENDICE)
    print(f"Gerado: {SAIDA_APENDICE}")
    print(f"  parágrafos: {len(b.doc.paragraphs)} | tabelas: {b.n_tabela} | figuras: {b.n_figura}")

    # 3. O Apêndice C — as capturas das execuções, com numeração C.x.
    c = Artigo(ORIGEM)
    apendice_c(c)
    c.salvar(SAIDA_APENDICE_C)
    print(f"Gerado: {SAIDA_APENDICE_C}")
    print(f"  parágrafos: {len(c.doc.paragraphs)} | tabelas: {c.n_tabela} | figuras: {c.n_figura}")

    # 4. A versão navegável do Apêndice C — é o alvo da URL citada nas Referências.
    print(f"Gerado: {gerar_apendice_c_md()}")

    # O artigo consome as imagens 1 a 3 da origem e o apêndice A as 4 a 13; juntos, todas
    # as 13. A verificação é sobre o CURSOR de imagens da origem, não sobre o total de
    # figuras: o corpo do artigo passou a numerar também as capturas de execução, lidas
    # de arquivo, e o Apêndice C é composto só por elas. Um desencontro no cursor
    # significaria legenda de diagrama trocada em silêncio — daí a verificação.
    if a.i_imagem != 3 or b.i_imagem != 13 or c.i_imagem != 0:
        raise SystemExit(
            f"Consumo inesperado das imagens da origem: artigo={a.i_imagem} (esperado 3), "
            f"apêndice A={b.i_imagem} (esperado 13), apêndice C={c.i_imagem} (esperado 0)."
        )
    if a.n_figura != 9 or b.n_figura != 10 or c.n_figura != len(CAPTURAS):
        raise SystemExit(
            f"Contagem inesperada de figuras: artigo={a.n_figura} (esperado 9), "
            f"apêndice A={b.n_figura} (esperado 10), "
            f"apêndice C={c.n_figura} (esperado {len(CAPTURAS)})."
        )


if __name__ == "__main__":
    main()

